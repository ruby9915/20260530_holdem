# -*- coding: utf-8 -*-
"""zca_vic 논문 초안을 MS Word(.docx)로 생성 — KCI 등재지 [별표1] 원고작성요령 준수.
실행: ../.venv/Scripts/python.exe build_docx.py   (zca_theory/ 에서)

규격(별표1):
- 용지 190x260mm, 여백 상하좌우 20, 머리말 15, 꼬리말 0
- 본문 2단(단간격 약 10), 기본 신명조 9pt, 줄간격 160%, 장평 95%, 자간 -5%
- 전치 순서: 국문 제목·저자·소속·요약·주제어 → 영문 제목·저자·소속·요약·주제어 → 본문 → 참고문헌 → 저자소개
- 목차 1./1.1/1.1.1 (아라비아), 표는 상단 중앙 <Table N>, 그림은 하단 <Figure N>
- 제목(1.) 12pt 진하게 중앙, 소문단(1.1) 10pt 진하게, 소소문단(1.1.1) 9.5pt 진하게
- 참고문헌은 본문 [n] 인용, 영문 작성
수식은 유니코드(μ ε ≈ ≠ → ∀ ∈ − ₁ ₂)로 표기.
신명조 폰트가 없으면 BODY_FONT 를 '바탕' 등으로 교체.
"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "신명조"      # 별표1 기본 폰트. 미설치 시 '바탕'으로 교체 가능.
JANGPYEONG = 95           # 장평 %
JAGAN = -5                # 자간 %


def _fmt_run(run, size, bold=False, italic=False, font=BODY_FONT,
             jang=JANGPYEONG, jagan=JAGAN):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    w = OxmlElement("w:w"); w.set(qn("w:val"), str(jang)); rPr.append(w)         # 장평
    val = round(jagan / 100.0 * size * 20)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(val)); rPr.append(sp)  # 자간


def para(doc, text="", size=9, bold=False, italic=False, align=AL.JUSTIFY,
         space_after=0, line=1.6, indent=0, hang=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(0)
    pf.line_spacing = line
    if indent:
        pf.left_indent = Pt(indent)
    if hang:
        pf.first_line_indent = Pt(-hang)
    if text:
        _fmt_run(p.add_run(text), size, bold, italic)
    return p


def h1(doc, text):   # 1. 서론 — 12pt 진하게 중앙, 문단아래 12pt
    p = para(doc, text, size=12, bold=True, align=AL.CENTER, space_after=12, line=1.3)
    p.paragraph_format.space_before = Pt(6)
    return p


def label(doc, text, size=9, align=AL.CENTER, bold=False, italic=False, after=0):
    return para(doc, text, size=size, bold=bold, italic=italic, align=align,
                space_after=after, line=1.3)


def table(doc, rows, caption):
    label(doc, caption, size=9, align=AL.CENTER, bold=True, after=2)   # 표 상단 중앙 <Table N>
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            cp = t.cell(i, j).paragraphs[0]
            cp.alignment = AL.CENTER
            cp.paragraph_format.line_spacing = 1.2
            _fmt_run(cp.add_run(str(cell)), 8.5, bold=(i == 0))
    para(doc, "", size=9, space_after=4)
    return t


# ════════════════════════════════════════════════════════════════════
doc = Document()
st = doc.styles["Normal"]
st.font.name = BODY_FONT
st.font.size = Pt(9)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
st.paragraph_format.line_spacing = 1.6


def setup_section(sec, cols):
    sec.page_width = Mm(190); sec.page_height = Mm(260)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(20)
    sec.left_margin = Mm(20); sec.right_margin = Mm(20)
    sec.header_distance = Mm(15); sec.footer_distance = Mm(0)
    sectPr = sec._sectPr
    c = sectPr.find(qn("w:cols"))
    if c is None:
        c = OxmlElement("w:cols"); sectPr.append(c)
    c.set(qn("w:num"), str(cols)); c.set(qn("w:space"), "360")


# ── 전치(1단): 국문 제목·저자·소속·요약·주제어 ──
label(doc, "비례배분 기여도 하의 영-기여 흡수(zero-credit absorption) 진단과", size=15, bold=True)
label(doc, "최소 가상비용 처방 — tabular Monte-Carlo 포커 학습의 형식적 사례연구", size=15, bold=True, after=8)
label(doc, "[저자명]*", size=10.5)
label(doc, "[소속기관] [직위]", size=9.5, after=8)
label(doc, "국문 요약", size=11, bold=True, after=4)
para(doc,
     "표 기반 Monte-Carlo 제어로 헤즈업 노리밋 홀덤을 학습할 때, 각 행동의 보상을 그 행동의 투자액 "
     "비율로 배분하는 비례배분(proportional) 기여도 정형화는 분산을 줄이는 합리적 선택이다. 본 연구는 "
     "이 방식이 비용 0 행동(예: CHECK)에 구조적 병리를 남김을 형식적으로 보인다. 비용 0 행동의 비례 "
     "credit은 매 에피소드 정확히 0이므로 그 Monte-Carlo 고정점은 행동의 참 가치와 무관하게 0에 "
     "고정되고, 모든 가용 행동이 음(−)의 기댓값인 노드에서 이 0은 열등한 비용 0 행동을 greedy 정책이 "
     "영구히 선택하게 만든다. 우리는 이를 영-기여 흡수(zero-credit absorption, ZCA)로 명명하고 최소 "
     "toy MDP에서 (i) 영-고정점, (ii) 흡수 정리, (iii) 낙관적 초기화의 일시적 흡수와의 질적 "
     "구별(영구적·고정점≠참값)을 증명한다. 또한 무시 가능한 가상비용(virtual information cost, VIC)이 "
     "흡수를 해소함을 보이고 그 임계 가상비용을 유도하며, 통제비교로 무정보 노이즈·동점 처리(tie-break)가 "
     "흡수를 탈출하지 못함을 보인다. 본 결과는 학습 seed·평가 표본과 무관한 고정점 사실이며, 실제 다단계 "
     "MDP의 성능 효과 및 분포 외(out-of-distribution) 일반화 인과는 검증되지 않은 향후 과제로 정직하게 "
     "분리한다.", size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("주제어 : "), 9, bold=True)
_fmt_run(p.add_run("강화학습, Monte-Carlo 제어, 기여도 배분, 보상 형성, 영-기여 흡수, tabular Q-learning, "
                   "헤즈업 노리밋 홀덤"), 9)

# 영문 제목·저자·소속·요약·주제어
label(doc, "Diagnosing Zero-Credit Absorption under Proportional Credit Assignment and a",
      size=12.5, bold=True)
label(doc, "Minimal Virtual-Cost Remedy: A Formal Case Study in Tabular Monte-Carlo Poker Learning",
      size=12.5, bold=True, after=8)
label(doc, "[Author Name]*", size=10.5)
label(doc, "[Affiliation], [Position]", size=9.5, after=8)
label(doc, "ABSTRACT", size=11, bold=True, after=4)
para(doc,
     "In tabular Monte-Carlo (MC) control for heads-up no-limit Texas hold'em, assigning each action "
     "a share of the terminal payoff proportional to its invested chips—proportional credit "
     "assignment—is a reasonable variance-reduction choice. We show formally that it introduces a "
     "structural pathology for zero-cost actions (e.g., CHECK): their proportional credit is exactly "
     "zero every episode, so the MC fixed point is pinned to zero regardless of true value, and at any "
     "node where all actions have negative expected value this zero makes a greedy policy permanently "
     "select the inferior zero-cost action. We name this zero-credit absorption (ZCA) and prove in a "
     "minimal toy MDP (i) the zero fixed point, (ii) an absorption theorem, and (iii) its qualitative "
     "distinction from the transient absorption of optimistic initialization (permanent; fixed point ≠ "
     "true value). We further show a negligible virtual information cost (VIC) dissolves the "
     "absorption, derive the exact threshold, and show by controlled comparison that uninformative "
     "noise and tie-breaking fail to escape it. These are fixed-point facts independent of seed and "
     "sample; the full-MDP performance effect and any link to out-of-distribution generalization are "
     "honestly separated as future work.", size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("Keywords : "), 9, bold=True)
_fmt_run(p.add_run("reinforcement learning, Monte-Carlo control, credit assignment, reward shaping, "
                   "zero-credit absorption, tabular Q-learning, heads-up no-limit hold'em"), 9)

# ── 본문(2단) ──
body = doc.add_section(WD_SECTION.CONTINUOUS)

h1(doc, "1. 서론")
para(doc, "표 기반 강화학습은 함수근사의 블랙박스 효과 없이 학습 동역학을 단일변수로 통제·관찰할 수 있다. "
          "본 연구의 무대는 헤즈업 노리밋 홀덤을 순수 tabular Monte-Carlo로 학습하는 환경이며, 목적은 강한 "
          "봇이 아니라 구조적 기여도 배분이 남기는 측정 가능한 병리의 진단이다.", space_after=6)
para(doc, "Monte-Carlo 제어에서 종단 보상을 행동에 배분하는 방식은 학습 신호의 분산·편향을 좌우한다. "
          "표준 Monte-Carlo는 편향이 없지만 분산이 크고, 자연스러운 대안은 각 행동이 판돈에 투입한 금액의 "
          "비율로 보상을 나누는 비례배분(proportional)이다. 그러나 본 연구는 이 합리적 선택이 비용 0 "
          "행동(투자액 0, 대표적으로 CHECK)에 구조적 함정을 내포함을 보인다. 비용 0 행동의 비례 credit은 "
          "분자가 0이라 항상 0이고, 그 가치 추정의 고정점은 실제 기댓값과 무관하게 0에 고정된다. 이 0은 "
          "음의 값으로 올바르게 학습된 다른 행동들보다 높아 보여, greedy 정책이 열등한 비용 0 행동을 "
          "영구히 선택한다.", space_after=6)
para(doc, "본 논문의 기여는 (1) 이 현상을 영-기여 흡수(ZCA)로 명명하고 toy MDP에서 형식적으로 증명하며, "
          "(2) ZCA가 낙관적 초기화와 고정점 수준에서 질적으로 구별됨을 보이고, (3) 무시 가능한 "
          "가상비용 VIC로 흡수를 해소하는 처방과 그 임계를 유도하며 무정보 대안과 통제비교하고, (4) 본 "
          "결과의 범위를 — 고정점 사실이되 성능·일반화 인과는 미검증임을 — 정직하게 분리한 것이다.",
     space_after=6)

h1(doc, "2. 관련 연구")
para(doc, "ZCA를 이루는 부품은 인접 문헌에 모두 존재하나, 본 연구는 그 부호를 뒤집어 같은 구조를 실패 "
          "모드로 진단한다. 협력게임 이론의 Shapley value[3] 및 Shapley Q-value[4]는 null-player "
          "공리(한계 기여 0 → 배분 0)를 공정성의 바람직한 성질로 둔다 — 본 연구의 \"비용 0 → credit 0\"은 "
          "그 정확한 대응물이나 진단의 방향이 반대다. 같은 직관은 difference rewards[1]·COMA[2]의 "
          "counterfactual 신용에도 깔려 있다. RUDDER[5]는 return-equivalent 재분배가 최적 정책을 보존함을 "
          "보이는데, 비례배분은 비용 0 행동에 대해 return-equivalent가 아니며 ZCA는 정확히 그 비등가 "
          "영역에서 발생한다(보존 정리의 대우).", space_after=6)
para(doc, "처방 측면에서 VIC는 potential-based reward shaping(PBRS)이 Q-value 초기화와 등가[6][7]인 틀 "
          "안에 위치한다. 다만 \"비용 0(no-op) 행동에 선택적·음·credit 기반 가상비용을 부여해 영-고정 "
          "흡수를 깬다\"는 처방을 명시한 선행은 확인되지 않았고, 의도상 가장 가까운 action-penalty[8]는 모든 "
          "행동에 일률적이라 선택적이지 않다. 한편 미방문/낙관적 0이 음수 행동을 일시적으로 흡수하는 현상은 "
          "잘 알려져 있고 비관적 초기화의 영구 미선택도 정식화되어 있으나[9], 이는 초기화·보상값 기인으로 "
          "비례 credit이 구조적으로 0을 생성하는 본 연구의 영-기여 흡수와 기제가 다르다. 단일 seed의 결과 "
          "오도 위험은 재현성 문헌[10][11][12]이 확립했으며, 본 연구는 이를 5장에서 부정적 결과로 보고한다.",
     space_after=6)

h1(doc, "3. 영-기여 흡수의 형식적 특성화")
para(doc, "문제 설정(toy MDP). 결정 상태 s에서 행동 a₁을 한 번 선택한다. 각 행동은 투자액 inv(a₁)≥0을 "
          "가지며 비용 0 행동(CHECK)은 inv=0이다. 선택 후 궤적은 투자액 c>0인 후속 행동을 적어도 하나 "
          "포함하고 종단 보상 P로 끝난다(γ=1, 종단 보상만). μ(a₁):=E[P|a₁]로 두면 참값은 q*(s,a₁)=μ(a₁)"
          "이다. 세 배분 방식의 행동별 return R(s,a₁)은 <Table 1>과 같고, MC 업데이트 Q←Q+α(R−Q)의 "
          "고정점은 E[R]이다(sample-average 또는 감소 α; 상수 α도 평균 보존).", space_after=6)
table(doc, [["방식 (scheme)", "R(s, a₁)"],
            ["표준 MC", "P"],
            ["비례배분 (PROP)", "[ inv(a₁) / (inv(a₁)+c) ] · P"],
            ["VIC", "PROP과 동일, 단 inv(CHECK) ← ε > 0"]],
      "<Table 1> 행동별 return R(s, a₁) / Per-action return")
para(doc, "Lemma 1 (표준 MC의 일치성). Q_std(s,a₁) → μ(a₁) = q*(s,a₁). 증명. R=P이므로 고정점은 "
          "E[P|a₁]=μ(a₁). ∎", space_after=4)
para(doc, "Lemma 2 (비용 0 행동의 영-고정점). Q_prop(s,a₁) = [inv(a₁)/(inv(a₁)+c)]·μ(a₁), 특히 "
          "Q_prop(s,CHECK) = (0/(0+c))·μ(CHECK) = 0  (∀ μ(CHECK)). 증명. CHECK는 분자가 0. ∎  CHECK의 "
          "고정점은 참값과 무관하게 0이며, R≡0이라 표본 분산도 0인 구조적 고정점(초기화 잔재 아님)이다. "
          "비용>0 행동은 μ(a₁)을 inv/(inv+c)<1로 수축하되 부호를 보존한다 — 이 비대칭이 ZCA의 핵심이다.",
     space_after=4)
para(doc, "Theorem (영-고정점에 의한 흡수). q*(CHECK)=μ_C < μ_B=q*(BET)이고 μ_B<0이면(BET이 낫지만 둘 "
          "다 −EV), 비례배분 greedy는 열등한 CHECK를, 표준 MC greedy는 최적 BET을 선택한다. 증명. "
          "Q_prop(CHECK)=0 > [b/(b+c)]·μ_B = Q_prop(BET)이나 μ_C<μ_B이므로 CHECK는 열등. 표준 MC는 "
          "Q=μ이므로 BET 선택. ∎", space_after=4)
para(doc, "흡수의 적용 범위(중요). 흡수는 덜 나쁜 대안조차 −EV일 때만 발생한다. 어떤 행동이 +EV이면 "
          "Q_prop>0=Q_prop(CHECK)라 정상 선택되므로, ZCA가 손상시키는 것은 \"모든 가용 행동이 −EV인 "
          "노드에서 덜 나쁜 것을 고르는 능력\"뿐이다.", space_after=4)
para(doc, "Proposition 1 (낙관적 초기화와의 구분). Q를 0으로 초기화한 표준 MC에서 미방문 "
          "행동의 0-선호는 충분한 방문 후 Q→μ로 소거되어 일시적이며 고정점은 참값이다(Lemma 1). "
          "대조적으로 Q_prop(CHECK)=0은 수렴 후에도 유지되는 고정점이며 참값과 다르다(Lemma 2). 즉 "
          "표면(\"0이 음수를 흡수\")은 같으나 낙관적 초기화는 \"아직 학습 못 해서\"(일시적, 고정점=참값), "
          "ZCA는 \"구조적으로 학습 불가라서\"(영구, 고정점=0≠참값)이다. 이 고정점의 차이가 본 진단의 "
          "핵심이다.", space_after=6)

h1(doc, "4. VIC: 임계 가상비용과 탈출 메커니즘 비교")
para(doc, "Proposition 2 (VIC 복원과 임계). CHECK에 가상 투자 ε>0을 주면 Q_vic(CHECK)=[ε/(ε+c)]·μ_C이다. "
          "Theorem 설정에서 greedy가 BET을 선택할 필요충분조건은 [b/(b+c)]·μ_B > [ε/(ε+c)]·μ_C이고, "
          "μ_C<0이므로 k := [b/(b+c)]·(μ_B/μ_C) ∈ (0,1)에 대해  ε > ε_min = k·c / (1−k).  ε이 임계를 "
          "넘으면 Q_vic(CHECK)가 음으로 내려가 복원되고, 너무 작으면 흡수가 잔존한다. ε_min은 일반적으로 "
          "작아 \"무시할 1칩이 임계만 넘으면 충분\"하다는 설계와 정합한다. 핵심은 VIC가 부여하는 값이 참값 "
          "μ_C에 연동(informed)된다는 점이다.", space_after=6)
para(doc, "탈출 메커니즘 통제비교(<Table 2>). 비례배분 위에 VIC, 무정보 노이즈(결정 시 N(0,σ²) 가산), "
          "tie-break(정확히 동률일 때만 CHECK 후순위), 고정 페널티(−κ)를 얹어 비교한다.", space_after=4)
table(doc, [["메커니즘 (mechanism)", "Q(CHECK)", "Q(BET)", "P(pick BET)", "판정"],
            ["표준 MC (대조)", "−5.0", "−1.0", "1.00", "ZCA 없음"],
            ["비례배분 (ZCA)", "0.0", "−0.5", "0.00", "흡수"],
            ["VIC (ε=1)", "−2.5", "−0.5", "1.00", "복원·informed"],
            ["노이즈 (σ=1)", "0.0", "−0.5", "0.36", "미복원"],
            ["tie-break", "0.0", "−0.5", "0.00", "미복원"],
            ["고정 페널티 (κ=1)", "−1.0", "−0.5", "1.00", "복원·무정보"]],
      "<Table 2> 탈출 메커니즘 비교 / Escape-mechanism comparison (μ_C=−5, μ_B=−1, b=c=1)")
para(doc, "두 가지가 갈린다. 첫째, strict 흡수 발동. ZCA는 0 > Q(BET) < 0인 부등식이지 동률이 아니므로, "
          "\"동률일 때만\" 작동하는 tie-break은 미발동(0≠−0.5)이며 흡수가 잔존한다. 노이즈도 고정점 0이 "
          "−0.5보다 높아 결정 시 잡음을 줘도 열등 CHECK를 약 64% 선택한다(학습 중 zero-mean 노이즈도 MC "
          "고정점을 안 바꿔 동일). 즉 tie-break·노이즈는 strict 흡수를 탈출하지 못한다. 둘째, informed "
          "여부. VIC와 고정 페널티는 복원하나, VIC의 고정점은 μ_C에 연동(informed)되어 서로 다른 μ_C의 "
          "다중 비용 0 상태에 자동 적응하는 반면 고정 페널티 −κ는 무정보다. 단 이 적응 이점의 성능 "
          "유의성은 실제 MDP 통제비교가 필요한 별도 문제다. 결국 VIC만 \"strict 흡수 해소 + 고정점 "
          "informed\"를 동시 충족한다.", space_after=6)

h1(doc, "5. 수치 검증과 한계")
para(doc, "검증. <Table 1>·<Table 2>의 모든 주장을 독립 Monte-Carlo 시뮬레이션(40만 에피소드)으로 "
          "검증하였고 결과는 <Table 3> 및 해석식과 일치한다. 특히 비례배분에서 Q(CHECK)=0.000 (≠ μ_C=−5)"
          "으로 흡수가, VIC에서 −2.50으로 복원이 재현된다. 임계도 ε_min=0.111에서 ε≤0.111→CHECK·"
          "ε>0.111→BET으로 소수점까지 일치한다. 검증 코드는 공개한다.", space_after=4)
table(doc, [["방식", "Q(CHECK)", "Q(BET)", "greedy", "판정"],
            ["표준 MC", "−4.99 (≈μ_C)", "−1.00", "BET", "정상(최적)"],
            ["비례배분", "0.000 (≠μ_C)", "−0.50", "CHECK", "흡수(열등)"],
            ["VIC", "−2.50", "−0.50", "BET", "복원"]],
      "<Table 3> 수치 검증 / Numerical verification (40만 ep, b=c=1, μ_C=−5, μ_B=−1, ε=1)")
para(doc, "범위와 한계(정직). (1) 본 결과는 학습 seed·평가 표본과 무관한 고정점 사실이다. 실제 2,048-셀 "
          "MDP에서는 c가 궤적마다 변하나 핵심(비용 0의 구조적 0-credit)은 동일하다. (2) 본 정리는 ZCA의 "
          "존재·특성화이지 VIC의 성능·일반화 인과가 아니다. 예비 실험에서 VIC 제거가 미학습 상대 대상 "
          "선택적 저하를 보였으나 학습 seed 전반의 통제실험에서 재현되지 않았다 — 따라서 \"VIC=분포 "
          "일반화의 필요조건\"이라는 강한 주장은 지지하지 않으며, 단일 seed 오도 가능성[10][11][12]과 "
          "정합하는 부정적 결과로 보고한다. 비용 0의 영-고정점(seed 무관)과 OOD 성능(seed 의존)은 별개 "
          "층위다. (3) VIC는 임계 이상이면 작동하는 한 처방이며 유일·최선이 아니다(4장의 고정 페널티 대비 "
          "성능 우열은 미검증). (4) ZCA의 부품은 알려져 있고 VIC는 PBRS≡Q-init[7] 틀 안에 있어 알고리즘 "
          "신규성은 제한적이다 — 기여는 알려진 구조의 미보고 실패 모드를 명명·형식화하고 임계를 유도하며 "
          "범위를 정직히 분리한 데 있다. (5) 단일 추상화·단일 게임에 한정된다.", space_after=6)

h1(doc, "6. 결론")
para(doc, "본 연구는 tabular Monte-Carlo 비례배분 기여도가 비용 0 행동에 남기는 구조적 병리를 영-기여 "
          "흡수(ZCA)로 명명하고, toy MDP에서 영-고정점과 흡수를 증명하였다. ZCA는 낙관적 초기화와 "
          "고정점 수준에서 질적으로 구별되며, 최소 가상비용 VIC가 명시적 임계 이상에서 "
          "이를 해소함을, 그리고 무정보 노이즈·tie-break은 strict 흡수를 탈출하지 못함을 보였다. 이 "
          "결과들은 seed·표본과 무관한 형식적 사실이며, 실제 다단계 MDP의 성능 효과와 분포 외 일반화 "
          "인과는 향후 과제로 분리하였다. 본 사례연구의 가치는 성능이 아니라 구조적 기여도 배분의 측정 "
          "가능한 병리를 형식적으로 진단하고 한계를 정직하게 드러낸 데 있다.", space_after=8)

para(doc, "참고문헌", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
refs = [
    '[1] D. H. Wolpert and K. Tumer, "Optimal payoff functions for members of collectives," Advances in Complex Systems, Vol. 4, No. 2/3, pp. 265-279, 2001.',
    '[2] J. Foerster, G. Farquhar, T. Afouras, N. Nardelli, and S. Whiteson, "Counterfactual multi-agent policy gradients," Proc. AAAI Conf. on Artificial Intelligence, pp. 2974-2982, 2018.',
    '[3] L. S. Shapley, "A value for n-person games," Contributions to the Theory of Games II, Princeton Univ. Press, pp. 307-317, 1953.',
    '[4] J. Wang, Y. Zhang, T.-K. Kim, and Y. Gu, "Shapley Q-value: A local reward approach to solve global reward games," Proc. AAAI Conf. on Artificial Intelligence, pp. 7285-7292, 2020.',
    '[5] J. A. Arjona-Medina, M. Gillhofer, M. Widrich, T. Unterthiner, J. Brandstetter, and S. Hochreiter, "RUDDER: Return decomposition for delayed rewards," Advances in Neural Information Processing Systems, pp. 13544-13555, 2019.',
    '[6] A. Y. Ng, D. Harada, and S. Russell, "Policy invariance under reward transformations: Theory and application to reward shaping," Proc. Int. Conf. on Machine Learning, pp. 278-287, 1999.',
    '[7] E. Wiewiora, "Potential-based shaping and Q-value initialization are equivalent," Journal of Artificial Intelligence Research, Vol. 19, pp. 205-208, 2003.',
    '[8] S. Koenig and R. G. Simmons, "The effect of representation and knowledge on goal-directed exploration with reinforcement-learning algorithms," Machine Learning, Vol. 22, pp. 227-250, 1996.',
    '[9] T. Rashid, B. Peng, W. Boehmer, and S. Whiteson, "Optimistic exploration even with a pessimistic initialisation," Proc. Int. Conf. on Learning Representations, 2020.',
    '[10] P. Henderson, R. Islam, P. Bachman, J. Pineau, D. Precup, and D. Meger, "Deep reinforcement learning that matters," Proc. AAAI Conf. on Artificial Intelligence, pp. 3207-3214, 2018.',
    '[11] C. Colas, O. Sigaud, and P.-Y. Oudeyer, "How many random seeds? Statistical power analysis in deep reinforcement learning experiments," arXiv:1806.08295, 2018.',
    '[12] R. Agarwal, M. Schwarzer, P. S. Castro, A. C. Courville, and M. G. Bellemare, "Deep reinforcement learning at the edge of the statistical precipice," Advances in Neural Information Processing Systems, pp. 29304-29320, 2021.',
    '[13] J. Kim, "PokerKit: A comprehensive Python library for fine-grained multi-variant poker game simulations," IEEE Trans. on Games, 2023.',
    '[14] R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.',
]
for r in refs:
    para(doc, r, size=9, align=AL.JUSTIFY, space_after=2, line=1.3, indent=16, hang=16)

para(doc, "저자소개", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
para(doc, "[저자명] (Author Name)  [정회원]", size=9, space_after=2)
para(doc, "[학사/석사/박사 취득연월 및 학위명]", size=9, space_after=2)
para(doc, "[현 소속기관 및 직위]  E-mail : [e-mail]", size=9, space_after=2)
para(doc, "관심분야 : 강화학습, 게임 AI, 학습 동역학 분석", size=9, space_after=2)

setup_section(doc.sections[0], cols=1)   # 전치 1단
setup_section(body, cols=2)              # 본문 2단

doc.save("zca_vic_논문초안.docx")
print("saved: zca_vic_논문초안.docx  (별표1 양식)")

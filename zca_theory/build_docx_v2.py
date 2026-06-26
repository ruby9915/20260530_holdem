# -*- coding: utf-8 -*-
"""2차 논문(v2) → MS Word(.docx). KCI [별표1] 양식. 측정 중심 3축(이론·실측·처방).
실행: ../.venv/Scripts/python.exe build_docx_v2.py   (zca_theory/ 에서)
소스: zca_vic_논문초안_v2.md. 수치는 results/ 검증값.
신명조 미설치 시 BODY_FONT='바탕'으로 교체."""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "신명조"
JANGPYEONG = 95
JAGAN = -5


def _fmt_run(run, size, bold=False, italic=False, font=BODY_FONT, jang=JANGPYEONG, jagan=JAGAN):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.name = font
    rPr = run._element.get_or_add_rPr(); rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font); rFonts.set(qn("w:hAnsi"), font); rFonts.set(qn("w:eastAsia"), font)
    w = OxmlElement("w:w"); w.set(qn("w:val"), str(jang)); rPr.append(w)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(round(jagan / 100.0 * size * 20))); rPr.append(sp)


def para(doc, text="", size=9, bold=False, italic=False, align=AL.JUSTIFY, space_after=0, line=1.6, indent=0, hang=0):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.alignment = align; pf.space_after = Pt(space_after); pf.space_before = Pt(0); pf.line_spacing = line
    if indent: pf.left_indent = Pt(indent)
    if hang: pf.first_line_indent = Pt(-hang)
    if text: _fmt_run(p.add_run(text), size, bold, italic)
    return p


def h1(doc, text):   # 1. 서론 — 12pt 진하게 중앙
    p = para(doc, text, size=12, bold=True, align=AL.CENTER, space_after=12, line=1.3)
    p.paragraph_format.space_before = Pt(6); return p


def h2(doc, text):   # 5.1 소문단 — 10pt 진하게 좌측
    p = para(doc, text, size=10, bold=True, align=AL.LEFT, space_after=4, line=1.3)
    p.paragraph_format.space_before = Pt(6); return p


def label(doc, text, size=9, align=AL.CENTER, bold=False, italic=False, after=0):
    return para(doc, text, size=size, bold=bold, italic=italic, align=align, space_after=after, line=1.3)


def table(doc, rows, caption):
    label(doc, caption, size=9, align=AL.CENTER, bold=True, after=2)
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            cp = t.cell(i, j).paragraphs[0]; cp.alignment = AL.CENTER; cp.paragraph_format.line_spacing = 1.15
            _fmt_run(cp.add_run(str(cell)), 8.5, bold=(i == 0))
    para(doc, "", size=9, space_after=4); return t


def setup_section(sec, cols):
    sec.page_width = Mm(190); sec.page_height = Mm(260)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(20); sec.left_margin = Mm(20); sec.right_margin = Mm(20)
    sec.header_distance = Mm(15); sec.footer_distance = Mm(0)
    sectPr = sec._sectPr; c = sectPr.find(qn("w:cols"))
    if c is None: c = OxmlElement("w:cols"); sectPr.append(c)
    c.set(qn("w:num"), str(cols)); c.set(qn("w:space"), "360")


# ════════════════════════════════════════════════════════════════════
doc = Document()
st = doc.styles["Normal"]; st.font.name = BODY_FONT; st.font.size = Pt(9)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT); st.paragraph_format.line_spacing = 1.6

# ── 전치(1단) ──
label(doc, "비례배분 기여도 하의 영-기여 흡수(zero-credit absorption):", size=15, bold=True)
label(doc, "형식적 진단, 실측, 그리고 최소 가상비용 처방", size=15, bold=True)
label(doc, "— tabular Monte-Carlo 헤즈업 홀덤 학습의 사례연구 —", size=11, bold=True, after=8)
label(doc, "[저자명]*", size=10.5)
label(doc, "[소속기관] [직위]", size=9.5, after=8)
label(doc, "국문 요약", size=11, bold=True, after=4)
para(doc,
     "표 기반 Monte-Carlo(MC) 제어로 헤즈업 노리밋 홀덤을 학습할 때, 각 행동의 보상을 그 행동의 투자액 "
     "비율로 배분하는 비례배분(proportional) 기여도 정형화는 분산을 줄이는 합리적 선택이다. 본 연구는 "
     "이 방식이 비용 0 행동(예: CHECK)에 구조적 병리를 남김을 (i) 형식적으로 증명하고, (ii) 실제 학습된 "
     "봇에서 측정하며, (iii) 최소 처방으로 완화한다. 이론적으로, 비용 0 행동의 비례 credit은 매 에피소드 "
     "정확히 0이므로 그 MC 고정점은 행동의 참값과 무관하게 0에 고정되고, 모든 가용 행동이 음(−)의 "
     "기댓값인 노드에서 이 0은 열등한 비용 0 행동을 greedy 정책이 영구히 선택하게 만든다. 우리는 이를 "
     "영-기여 흡수(zero-credit absorption, ZCA)로 명명하고 최소 toy MDP에서 증명하며, 이것이 낙관적 "
     "초기화의 일시적 흡수와 질적으로 구별됨(영구적·고정점≠참값)을 보인다. 실측에서, 2,048-셀 tabular "
     "봇의 Q(CHECK)는 표준 MC가 −73~+120칩으로 펼쳐지는 데 반해 비례배분에서는 [−0.5, +1.4]칩의 "
     "면도날 띠로 붕괴(평균 |Q(CHECK)| 4.48 → 0.065, 약 69배)하여 영-고정점이 실제로 관측된다. "
     "가상비용(VIC)은 이 핀을 푸는 것이 아니라(고정 비율 약 99% 유지) CHECK의 argmax 흡수율을 "
     "25.4%→7.8%로 완화하며, 100k×5 정밀평가에서 미학습(OOD) 상대 일반화를 선택적으로 좌우한다. 끝으로, "
     "단일 학습 seed의 한계와 학습-seed 전반에서 일부 효과가 재현되지 않음을 부정적 결과로 정직하게 "
     "보고한다.", size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("주제어 : "), 9, bold=True)
_fmt_run(p.add_run("강화학습, Monte-Carlo 제어, 기여도 배분, 보상 형성, 영-기여 흡수, 분포 외 일반화, "
                   "tabular Q-learning, 헤즈업 노리밋 홀덤"), 9)

label(doc, "Zero-Credit Absorption under Proportional Credit Assignment: Formal Diagnosis,",
      size=12, bold=True)
label(doc, "Measurement, and a Minimal Virtual-Cost Remedy — A Case Study in Tabular MC Hold'em",
      size=12, bold=True, after=8)
label(doc, "[Author Name]*", size=10.5)
label(doc, "[Affiliation], [Position]", size=9.5, after=8)
label(doc, "ABSTRACT", size=11, bold=True, after=4)
para(doc,
     "In tabular Monte-Carlo (MC) control for heads-up no-limit Texas hold'em, assigning each action "
     "a share of the terminal payoff proportional to its invested chips—proportional credit "
     "assignment—is a reasonable variance-reduction choice. We show it leaves a structural pathology "
     "for zero-cost actions (e.g., CHECK) by (i) proving it, (ii) measuring it in a trained agent, and "
     "(iii) mitigating it with a minimal remedy. Theoretically, a zero-cost action's proportional "
     "credit is exactly zero every episode, so its MC fixed point is pinned to zero regardless of true "
     "value; where all actions have negative expected value, this zero makes a greedy policy "
     "permanently select the inferior zero-cost action. We name this Zero-Credit Absorption (ZCA), "
     "prove it in a minimal toy MDP, and show it is distinct from the transient absorption of "
     "optimistic initialization (permanent; fixed point ≠ true value). Empirically, in a 2,048-cell "
     "agent the learned Q(CHECK) spans −73 to +120 chips under standard MC but collapses to a razor "
     "band of [−0.5, +1.4] chips under proportional credit (mean |Q(CHECK)| 4.48 → 0.065, ≈69×). A "
     "Virtual Information Cost (VIC) does not unpin this value (≈99% remains pinned) but reduces "
     "CHECK's argmax-absorption from 25.4% to 7.8%, and in 100k×5 evaluation it selectively governs "
     "generalization to unseen (OOD) opponents. We honestly report, as a negative result, the "
     "single-training-seed limitation and the non-reproducibility of some effects across seeds.",
     size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("Keywords : "), 9, bold=True)
_fmt_run(p.add_run("reinforcement learning, Monte-Carlo control, credit assignment, reward shaping, "
                   "zero-credit absorption, out-of-distribution generalization, tabular Q-learning, "
                   "heads-up no-limit hold'em"), 9)

# ── 본문(2단) ──
body = doc.add_section(WD_SECTION.CONTINUOUS)

h1(doc, "1. 서론")
para(doc, "표 기반 강화학습은 함수근사의 블랙박스 효과 없이 학습 동역학을 단일변수로 통제·관찰할 수 있어, "
          "\"무엇이 학습 안정성과 분포 강건성을 결정하는가\"를 규명하는 데 유용한 실험 환경을 제공한다. 본 "
          "연구의 무대는 헤즈업 노리밋 홀덤을 순수 tabular Monte-Carlo로 학습하는 환경이며, 목적은 강한 "
          "봇이 아니라 구조적 기여도 배분이 남기는 측정 가능한 병리의 진단이다.", space_after=6)
para(doc, "Monte-Carlo 제어에서 종단 보상을 행동에 배분하는 방식은 학습 신호의 분산·편향을 좌우한다. "
          "표준 Monte-Carlo는 편향이 없지만 분산이 크고, 자연스러운 대안은 각 행동이 판돈에 투입한 금액의 "
          "비율로 보상을 나누는 비례배분(proportional)이다. 그러나 본 연구는 이 합리적 선택이 비용 0 "
          "행동(투자액 0, 대표적으로 CHECK)에 구조적 함정을 내포함을 보인다. 비용 0 행동의 비례 credit은 "
          "분자가 0이라 항상 0이고, 그 가치 추정의 고정점은 실제 기댓값과 무관하게 0에 고정된다. 이 0은 "
          "음의 값으로 올바르게 학습된 다른 행동들보다 높아 보여, greedy 정책이 열등한 비용 0 행동을 "
          "영구히 선택한다.", space_after=6)
para(doc, "본 논문의 기여는 세 축으로 구성된다. (1) 이론 — 위 현상을 영-기여 흡수(ZCA)로 명명하고 최소 "
          "toy MDP에서 영-고정점과 흡수를 형식적으로 증명하며(3장), 낙관적 초기화와 고정점 수준에서 질적으로 "
          "구별됨을 보인다. (2) 처방 — 비용 0 행동에 무시 가능한 가상비용(VIC)을 부여하는 최소 개입과 그 "
          "임계를 유도하고, 무정보 노이즈·동점 처리가 본 흡수의 등가물이 아님을 보인다(4장). (3) 실측 — 실제 "
          "2,048-셀 tabular 봇에서 ZCA를 측정한다(5장): Q(CHECK)의 영-고정점(약 69배 붕괴), VIC의 흡수 "
          "완화(25.4%→7.8%), 미학습(OOD) 상대 일반화와의 연관, 학습-seed 취약성의 정직한 부정적 결과.",
     space_after=6)

h1(doc, "2. 관련 연구")
para(doc, "ZCA를 이루는 부품은 인접 문헌에 모두 존재하나, 본 연구는 그 부호를 뒤집어 같은 구조를 실패 "
          "모드로 진단한다. 협력게임 이론의 Shapley value[3] 및 Shapley Q-value[4]는 null-player "
          "공리(한계 기여 0 → 배분 0)를 공정성의 바람직한 성질로 둔다 — 본 연구의 \"비용 0 → credit 0\"은 "
          "그 정확한 대응물이나 진단의 방향이 반대다. 같은 직관은 difference rewards[1]·COMA[2]에도 깔려 "
          "있다. RUDDER[5]는 return-equivalent 재분배가 최적 정책을 보존함을 보이는데, 비례배분은 비용 0 "
          "행동에 대해 return-equivalent가 아니며 ZCA는 정확히 그 비등가 영역에서 발생한다(보존 정리의 "
          "대우).", space_after=6)
para(doc, "처방 측면에서 VIC는 potential-based reward shaping(PBRS)이 Q-value 초기화와 등가[6][7]인 틀 "
          "안에 위치한다. 다만 \"비용 0(no-op) 행동에 선택적·음·credit 기반 가상비용을 부여해 영-고정 "
          "흡수를 깬다\"는 처방을 명시한 선행은 확인되지 않았고, 의도상 가장 가까운 action-penalty[8]는 모든 "
          "행동에 일률적이라 선택적이지 않다. 미방문/낙관적 0이 음수 행동을 일시적으로 흡수하는 현상은 잘 "
          "알려져 있고 비관적 초기화의 영구 미선택도 정식화되어 있으나[9], 이는 초기화·보상값 기인으로 "
          "비례 credit이 구조적으로 0을 생성하는 본 연구의 영-기여 흡수와 기제가 다르다. 단일 seed의 결과 "
          "오도 위험은 재현성 문헌[10][11][12]이 확립했으며, 본 연구는 이를 5장에서 부정적 결과로 보고한다.",
     space_after=6)

h1(doc, "3. 영-기여 흡수의 형식적 특성화")
para(doc, "문제 설정(toy MDP). 결정 상태 s에서 행동 a₁을 한 번 선택한다. 각 행동은 투자액 inv(a₁)≥0을 "
          "가지며 비용 0 행동(CHECK)은 inv=0이다. 선택 후 궤적은 투자액 c>0인 후속 행동을 적어도 하나 "
          "포함하고 종단 보상 P로 끝난다(γ=1, 종단 보상만). μ(a₁):=E[P|a₁]로 두면 참값은 q*(s,a₁)=μ(a₁)"
          "이다. 세 배분 방식의 행동별 return은 <Table 1>과 같고, MC 업데이트 Q←Q+α(R−Q)의 고정점은 "
          "E[R]이다.", space_after=6)
table(doc, [["방식 (scheme)", "R(s, a₁)"],
            ["표준 MC", "P"],
            ["비례배분 (PROP)", "[ inv(a₁) / (inv(a₁)+c) ] · P"],
            ["VIC", "PROP과 동일, 단 inv(CHECK) ← ε > 0"]],
      "<Table 1> 행동별 return R(s, a₁) / Per-action return")
para(doc, "Lemma 1 (표준 MC의 일치성). Q_std(s,a₁) → μ(a₁) = q*(s,a₁). 증명. R=P이므로 고정점은 "
          "E[P|a₁]=μ(a₁). ∎", space_after=4)
para(doc, "Lemma 2 (비용 0 행동의 영-고정점). Q_prop(s,CHECK)=(0/(0+c))·μ(CHECK)=0 (∀ μ(CHECK)). 증명. "
          "CHECK는 분자가 0. ∎  CHECK의 고정점은 참값과 무관하게 0이며, R≡0이라 표본 분산도 0인 구조적 "
          "고정점(초기화 잔재 아님)이다. 비용>0 행동은 μ를 inv/(inv+c)<1로 수축하되 부호를 보존한다 — 이 "
          "비대칭이 ZCA의 핵심이다.", space_after=4)
para(doc, "Theorem (영-고정점에 의한 흡수). q*(CHECK)=μ_C < μ_B=q*(BET)이고 μ_B<0이면(BET이 낫지만 둘 "
          "다 −EV), 비례배분 greedy는 열등한 CHECK를, 표준 MC greedy는 최적 BET을 선택한다. 증명. "
          "Q_prop(CHECK)=0 > [b/(b+c)]·μ_B=Q_prop(BET)이나 μ_C<μ_B이므로 CHECK는 열등. 표준 MC는 Q=μ"
          "이므로 BET 선택. ∎  적용 범위(중요): 흡수는 덜 나쁜 대안조차 −EV일 때만 발생한다 — 어떤 행동이 "
          "+EV이면 정상 선택되므로, ZCA가 손상시키는 것은 \"모든 가용 행동이 −EV인 노드에서 덜 나쁜 것을 "
          "고르는 능력\"뿐이다.", space_after=4)
para(doc, "Proposition 1 (낙관적 초기화와의 구분). Q를 0으로 초기화한 표준 MC에서 미방문 행동의 0-선호는 "
          "충분한 방문 후 Q→μ로 소거되어 일시적이며 고정점은 참값이다(Lemma 1). 대조적으로 "
          "Q_prop(CHECK)=0은 수렴 후에도 유지되는 고정점이며 참값과 다르다(Lemma 2). 즉 표면(\"0이 음수를 "
          "흡수\")은 같으나 낙관적 초기화는 \"아직 학습 못 해서\"(일시적, 고정점=참값), ZCA는 \"구조적으로 "
          "학습 불가라서\"(영구, 고정점=0≠참값)이다.", space_after=4)
para(doc, "확장(향후 과제). 후속 투자 c를 궤적마다 변하는 확률변수로, 결정점을 다단계로 일반화해도 비용 0 "
          "행동의 credit은 매 실현에서 0이므로 영-고정점은 유지된다. 다단계에서는 CHECK의 참 연속가치가 클 "
          "수 있어(팟 컨트롤·트랩 등) 핀 오차가 더 커지며, 이는 5장에서 관측된 정책의 수동화와 정합한다.",
     space_after=6)

h1(doc, "4. VIC: 임계 가상비용과 탈출 메커니즘 비교")
para(doc, "Proposition 2 (VIC 복원과 임계). CHECK에 가상 투자 ε>0을 주면 Q_vic(CHECK)=[ε/(ε+c)]·μ_C이다. "
          "Theorem 설정에서 greedy가 BET을 선택할 필요충분조건은 [b/(b+c)]·μ_B > [ε/(ε+c)]·μ_C이고, "
          "μ_C<0이므로 k := [b/(b+c)]·(μ_B/μ_C) ∈ (0,1)에 대해  ε > ε_min = k·c / (1−k).  ε이 임계를 "
          "넘으면 Q_vic(CHECK)가 음으로 내려가 복원된다. 핵심은 VIC가 부여하는 값이 참값 μ_C에 "
          "연동(informed)된다는 점이다.", space_after=6)
para(doc, "탈출 메커니즘 통제비교(<Table 2>). 비례배분 위에 VIC, 무정보 노이즈(결정 시 N(0,σ²) 가산), "
          "tie-break(정확히 동률일 때만 CHECK 후순위), 고정 페널티(−κ)를 얹어 비교한다.", space_after=4)
table(doc, [["메커니즘", "Q(CHECK)", "Q(BET)", "P(pick BET)", "판정"],
            ["표준 MC (대조)", "−5.0", "−1.0", "1.00", "ZCA 없음"],
            ["비례배분 (ZCA)", "0.0", "−0.5", "0.00", "흡수"],
            ["VIC (ε=1)", "−2.5", "−0.5", "1.00", "복원·informed"],
            ["노이즈 (σ=1)", "0.0", "−0.5", "0.36", "미복원"],
            ["tie-break", "0.0", "−0.5", "0.00", "미복원"],
            ["고정 페널티 (κ=1)", "−1.0", "−0.5", "1.00", "복원·무정보"]],
      "<Table 2> 탈출 메커니즘 비교 / Escape-mechanism comparison (μ_C=−5, μ_B=−1, b=c=1)")
para(doc, "ZCA의 흡수는 0 > Q(BET) < 0인 부등식이지 동률이 아니므로 \"동률일 때만\" 작동하는 tie-break은 "
          "미발동(0≠−0.5)이며, 노이즈도 고정점 0이 −0.5보다 높아 결정 시 잡음을 주어도 열등 CHECK를 약 "
          "64% 선택한다. 즉 tie-break·노이즈는 strict 흡수를 탈출하지 못한다. VIC와 고정 페널티만 복원하나, "
          "VIC의 고정점만이 참값 μ_C에 연동되어(informed) 서로 다른 μ_C의 다중 비용 0 상태에 자동 "
          "적응한다.", space_after=6)

h1(doc, "5. 실제 MDP에서의 ZCA 실측")
para(doc, "본 장은 toy 이론이 실제 학습 시스템에서 측정 가능한 병리로 나타남을 보인다. 모든 수치는 학습된 "
          "Q-table에 대한 직접 진단(analyze_qcheck.py) 또는 100k×5 정밀평가에서 얻었다.", space_after=4)

h2(doc, "5.1 실험 환경")
para(doc, "헤즈업 노리밋 홀덤을 (라운드 × 포지션 × 핸드 버킷 × 직전 행동)의 2,048-셀 추상화 위에서 8개 "
          "행동의 tabular Q-table로 학습한다. 탐색은 softmax(온도 감쇠), 학습량은 2M 에피소드, 학습 상대는 "
          "5종 성향 페르소나(TAG/LAG/Maniac/Station/Nit)이다. 평가 단위는 mbb/g(=payoff×500)이며, 결론은 "
          "모두 100k 게임 × 5회(평가 seed 1000–1004) 정밀평가에서만 인용한다(200게임 체크포인트는 미학습 "
          "상대에서 잡음 지배). 학습 상대가 페르소나이므로 vs Random은 한 번도 학습하지 않은 미학습 "
          "분포(OOD), vs TAG는 학습 분포 내(ID)이다.", space_after=4)

h2(doc, "5.2 ZCA 지문 — Q(CHECK)의 영-고정점 측정")
para(doc, "표준 MC(PURE)와 비례배분(PROP)으로 학습한 Q-table에서 비용 0 행동 CHECK의 가치 분포를 "
          "측정한다(<Table 3>). 진단 지표: mean|Q(CHECK)|(영-고정이면 약 0), pinned(|Q|<1칩 셀 비율), "
          "learned(|Q|≥1칩 비율), range(최소~최대).", space_after=4)
table(doc, [["기여도 배분", "mean|Q(CHECK)|", "Q(CHECK) 범위(칩)", "pinned(|Q|<1)", "learned(|Q|≥1)"],
            ["표준 MC (PURE)", "4.48", "[−72.9, +119.8]", "72.6%", "27.4%"],
            ["비례배분 (PROP, VIC-off)", "0.065", "[−0.48, +1.39]", "99.5%", "0.5%"],
            ["비례배분 (PROP, VIC-on)", "0.121", "[−0.94, +1.06]", "99.4%", "0.6%"]],
      "<Table 3> 학습된 Q(CHECK)의 분포 / Distribution of learned Q(CHECK) (active 셀, 직접 진단)")
para(doc, "표준 MC에서 Q(CHECK)는 −73~+120칩으로 넓게 펼쳐지고 27.4%의 셀이 0에서 유의하게 벗어나 "
          "학습되는 반면, 비례배분에서는 평균 절댓값이 4.48 → 0.065로 약 69배 붕괴하고 99.5%의 셀이 ±1칩 "
          "이내에 핀되며 0.5%만이 학습된다. 즉 Lemma 2의 영-고정점이 실제 2,048-셀 시스템에서 면도날 "
          "띠로 직접 관측된다. 이 붕괴는 학습 seed·평가 표본과 무관한 구조적 측정이다.", space_after=4)

h2(doc, "5.3 흡수와 VIC의 효과 — 핀이 아니라 argmax 지배의 완화")
para(doc, "ZCA의 해악은 0 자체가 아니라 그 0이 음(−)으로 학습된 행동을 greedy argmax에서 흡수하는 데 "
          "있다. 이를 zca_dominance(CHECK가 argmax이면서 |Q(CHECK)|<1칩이고 다른 행동에 Q<−1칩이 존재하는 "
          "셀 비율)로 측정한다(<Table 4>).", space_after=4)
table(doc, [["지표", "VIC-off", "VIC-on"],
            ["Q(CHECK) pinned (|Q|<1)", "99.5%", "99.4%"],
            ["zca_dominance (흡수)", "25.4%", "7.8%"]],
      "<Table 4> VIC의 흡수 완화 효과 / VIC's mitigation of absorption (PROP, mixed)")
para(doc, "결정적으로, VIC는 Q(CHECK)의 핀을 풀지 않는다(VIC-on도 약 99% pinned). 변하는 것은 흡수율로, "
          "CHECK가 음수 행동을 지배하는 비율이 25.4% → 7.8%로 약 3.3배 완화된다. 즉 VIC는 \"참값을 0에서 "
          "멀리 옮기는\" 처방이 아니라, 학습된 음수 행동의 영구 배제를 막는(argmax 흡수만 미세 음으로 "
          "해소하는) 처방이다. 진짜 미학습 행동의 0(탐색 가치)은 보존된다.", space_after=4)

h2(doc, "5.4 OOD/ID 해리 — VIC 인과 ablation (100k×5)")
para(doc, "학습 방식(single/cycle/mixed) × VIC(on/off)의 2×3 통제실험에서 VIC만 단일 변경하고 100k×5로 "
          "평가한다(<Table 5>, 전 런 학습 seed=42).", space_after=4)
table(doc, [["방식", "VIC", "vs Random (OOD)", "vs TAG (ID)", "판정"],
            ["single", "on", "+987.1 (76.7)", "+940.4 (16.5)", "동시 흑자"],
            ["single", "off", "−3339.4 (143.8)", "+795.9 (15.8)", "OOD 붕괴"],
            ["cycle", "on", "+681.3 (51.4)", "+868.3 (12.1)", "동시 흑자"],
            ["cycle", "off", "−346.5 (17.8)", "+814.8 (4.4)", "OOD 붕괴"],
            ["mixed", "on", "+1083.9 (110.5)", "+885.2 (10.9)", "동시 흑자"],
            ["mixed", "off", "−1261.8 (22.5)", "+791.2 (20.3)", "OOD 붕괴"]],
      "<Table 5> VIC 2×3 인과 ablation (mbb/g, 괄호=회차 SD)")
para(doc, "세 방식 모두 VIC를 끄면 vs Random(OOD)이 큰 적자로 붕괴(−346~−3,339)하지만 vs TAG(ID)는 "
          "+791~+815로 유지된다. 즉 ZCA 흡수는 학습 분포 내 성능은 보존하면서 미학습 상대로의 일반화만 "
          "선택적으로 파괴한다. 이 부호반전 해리 자체가 본 연구의 진단 기여다.", space_after=4)

h2(doc, "5.5 표준 MC로의 우회는 해가 아니다 (Pure vs Prop)")
para(doc, "\"애초에 표준 MC를 썼으면 ZCA가 없지 않냐\"는 반박을 직접 검정한다. mixed 학습과 100% "
          "매칭하고 기여도 배분만 PURE로 바꾼 단일 변경의 100k×5 결과(<Table 6>).", space_after=4)
table(doc, [["기여도 배분", "vs Random", "vs TAG", "판정"],
            ["PURE", "+12.0 (38.8)", "+76.4 (8.8)", "약·브레이크이븐"],
            ["PROP + VIC-on", "+1083.9 (110.5)", "+885.2 (10.9)", "동시 흑자"],
            ["PROP + VIC-off", "−1261.8 (22.5)", "+791.2 (20.3)", "OOD 붕괴"]],
      "<Table 6> Pure vs Prop (mixed, 100k×5, mbb/g)")
para(doc, "PURE는 ZCA가 없어도(<Table 3>처럼 Q(CHECK)가 ±100칩 분포) 어디서나 약하다(+12/+76, PROP+VIC "
          "대비 vs TAG는 약 12배, vs Random은 그보다 훨씬 큰 폭으로 열위). 따라서 VIC의 가치는 \"ZCA "
          "제거\" 자체가 아니라 저분산 비례 기여도를 보존한 채 고정점 흡수만 해제하는 데 있다.", space_after=4)

h2(doc, "5.6 정직한 부정적 결과 — 학습-seed 취약성")
para(doc, "<Table 5>의 깨끗한 해리는 전 런이 학습 seed=42 단일이라는 한계를 갖는다(회차 SD는 평가 "
          "분산이지 학습 분산이 아니다). 학습 seed를 1–5로 바꾼 sweep에서, ID(vs TAG) 보존은 전 seed에서 "
          "유지되나 OOD(vs Random) 성패는 ±수백~수천 mbb/g로 요동했다 — cycle/mixed의 VIC-on은 5개 중 "
          "3개에서 음(−)으로, single의 VIC-off는 5개 중 4개에서 흑자로 나타나 seed=42의 극단값이 대표값이 "
          "아님이 드러났다. 따라서 \"VIC가 분포 일반화의 필요조건\"이라는 강한 인과 주장은 본 연구가 "
          "지지하지 않으며, 이는 단일 seed가 결과를 오도할 수 있다는 재현성 문헌[10][11][12]과 정합하는 "
          "부정적 결과다. 단, 5.2–5.3의 구조적 측정(영-고정점, 흡수율 완화)은 학습 seed와 무관한 사실이며, "
          "seed 의존적인 것은 5.4의 OOD 성능 크기다 — 양자는 별개 층위다.", space_after=6)

h1(doc, "6. 논의와 한계")
para(doc, "본 연구는 증명하는 것과 증명하지 않는 것을 분리한다. (1) 방어 가능(robust): 비례배분이 비용 0 "
          "행동에 영-고정점을 남긴다는 이론(3장)과 그 실측(5.2, 약 69배 붕괴)은 학습 seed·표본과 무관하다. "
          "VIC가 흡수율을 완화한다는 측정(5.3, 25.4%→7.8%)도 구조적이다. (2) 약하거나 미검증: OOD 일반화 "
          "성능의 크기와 그 인과(5.4)는 학습 seed에 의존하며(5.6), VIC의 성능 우월성이 고정 페널티 등 "
          "무정보 처방 대비 유의한지는 실제 MDP의 seed-sweep 통제비교가 필요하다. (3) 신규성의 경계: ZCA의 "
          "부품(null-player 공리·return-equivalence·낙관적 초기화·PBRS≡Q-init)은 모두 알려져 있고 VIC는 "
          "PBRS≡Q-init 등가[7] 틀 안에 있어 알고리즘 신규성은 제한적이다 — 기여는 알려진 구조의 미보고 "
          "실패 모드를 명명·형식화하고 실제 시스템에서 측정하며 범위를 정직히 분리한 데 있다. (4) 단일 "
          "추상화·단일 게임에 한정된다.", space_after=6)

h1(doc, "7. 결론")
para(doc, "본 연구는 tabular Monte-Carlo 비례배분 기여도가 비용 0 행동에 남기는 구조적 병리를 영-기여 "
          "흡수(ZCA)로 명명하고, 이를 세 축에서 닫았다 — toy MDP에서 영-고정점과 흡수를 증명하고(낙관적 "
          "초기화와 고정점 수준에서 구별), 실제 2,048-셀 봇에서 Q(CHECK)의 약 69배 붕괴와 흡수율을 "
          "측정하며, 최소 가상비용 VIC가 그 흡수를 25.4%→7.8%로 완화하고 미학습 상대 일반화를 선택적으로 "
          "좌우함을 보였다. 동시에 학습-seed 취약성으로 강한 인과 주장이 재현되지 않음을 부정적 결과로 "
          "정직하게 보고하였다. 본 사례연구의 가치는 성능이 아니라, 구조적 기여도 배분의 병리를 증명·측정·"
          "완화·정직한 한계 보고로 일관되게 다룬 데 있다.", space_after=8)

para(doc, "참고문헌", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
refs = [
    '[1] D. H. Wolpert and K. Tumer, "Optimal payoff functions for members of collectives," Advances in Complex Systems, Vol. 4, No. 2/3, pp. 265-279, 2001.',
    '[2] J. Foerster, G. Farquhar, T. Afouras, N. Nardelli, and S. Whiteson, "Counterfactual multi-agent policy gradients," Proc. AAAI Conf. on Artificial Intelligence, pp. 2974-2982, 2018.',
    '[3] L. S. Shapley, "A value for n-person games," Contributions to the Theory of Games II, Princeton Univ. Press, pp. 307-317, 1953.',
    '[4] J. Wang, Y. Zhang, T.-K. Kim, and Y. Gu, "Shapley Q-value: A local reward approach to solve global reward games," Proc. AAAI Conf. on Artificial Intelligence, pp. 7285-7292, 2020.',
    '[5] J. A. Arjona-Medina, M. Gillhofer, M. Widrich, T. Unterthiner, J. Brandstetter, and S. Hochreiter, "RUDDER: Return decomposition for delayed rewards," Advances in Neural Information Processing Systems, pp. 13544-13555, 2019.',
    '[6] A. Y. Ng, D. Harada, and S. Russell, "Policy invariance under reward transformations: Theory and application to reward shaping," Proc. Int. Conf. on Machine Learning, pp. 278-287, 1999.',
    '[7] E. Wiewiora, "Potential-based shaping and Q-value initialization are equivalent," Journal of Artificial Intelligence Research, Vol. 19, pp. 205-208, 2003.',
    '[8] S. Koenig and R. G. Simmons, "The effect of representation and knowledge on goal-directed exploration with reinforcement-learning algorithms," Machine Learning, Vol. 22, pp. 227-250, 1996.',
    '[9] T. Rashid, B. Peng, W. Boehmer, and S. Whiteson, "Optimistic exploration even with a pessimistic initialisation," Proc. Int. Conf. on Learning Representations, 2020.',
    '[10] P. Henderson, R. Islam, P. Bachman, J. Pineau, D. Precup, and D. Meger, "Deep reinforcement learning that matters," Proc. AAAI Conf. on Artificial Intelligence, pp. 3207-3214, 2018.',
    '[11] C. Colas, O. Sigaud, and P.-Y. Oudeyer, "How many random seeds? Statistical power analysis in deep reinforcement learning experiments," arXiv:1806.08295, 2018.',
    '[12] R. Agarwal, M. Schwarzer, P. S. Castro, A. C. Courville, and M. G. Bellemare, "Deep reinforcement learning at the edge of the statistical precipice," Advances in Neural Information Processing Systems, pp. 29304-29320, 2021.',
    '[13] J. Kim, "PokerKit: A comprehensive Python library for fine-grained multi-variant poker game simulations," IEEE Trans. on Games, 2023.',
    '[14] R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.',
]
for r in refs:
    para(doc, r, size=9, align=AL.JUSTIFY, space_after=2, line=1.3, indent=16, hang=16)

para(doc, "저자소개", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
para(doc, "[저자명] (Author Name)  [정회원]", size=9, space_after=2)
para(doc, "[학사/석사/박사 취득연월 및 학위명]", size=9, space_after=2)
para(doc, "[현 소속기관 및 직위]  E-mail : [e-mail]", size=9, space_after=2)
para(doc, "관심분야 : 강화학습, 게임 AI, 학습 동역학 분석", size=9, space_after=2)

setup_section(doc.sections[0], cols=1)
setup_section(body, cols=2)
doc.save("zca_vic_논문초안_v2.docx")
print("saved: zca_vic_논문초안_v2.docx  (별표1 양식, v2 3축)")

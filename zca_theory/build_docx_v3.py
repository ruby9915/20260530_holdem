# -*- coding: utf-8 -*-
"""3차 논문(v3.1) → MS Word(.docx). KCI [별표1] 양식. 소스: zca_vic_논문초안_v3.md (2026-07-06 재프레임).
v3.1 반영: 서론 재프레임(검증환경 선택 4사유) · Lemma 3(계열) · 5.3 전이구간(E7) · 5.3.1 3자 비교 ·
5.5 경쟁 처방 비교(E8, Table 4) · VII 향후 연구(V2G·IoT) · 참고문헌 [15]~[18] · 그림 3장(말미 1단).
실행: ../.venv/Scripts/python.exe build_docx_v3.py   (zca_theory/ 에서)
신명조 미설치 시 BODY_FONT='바탕'."""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = "신명조"
JANGPYEONG = 95
JAGAN = -5


def _fmt_run(run, size, bold=False, italic=False, font=BODY_FONT, jang=JANGPYEONG, jagan=JAGAN):
    run.font.size = Pt(size); run.font.bold = bold; run.font.italic = italic; run.font.name = font
    rPr = run._element.get_or_add_rPr(); rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font); rFonts.set(qn("w:hAnsi"), font); rFonts.set(qn("w:eastAsia"), font)
    w = OxmlElement("w:w"); w.set(qn("w:val"), str(jang)); rPr.append(w)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(round(jagan / 100.0 * size * 20))); rPr.append(sp)


def para(doc, text="", size=9, bold=False, italic=False, align=AL.JUSTIFY, space_after=0, line=1.6, indent=0, hang=0):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.alignment = align; pf.space_after = Pt(space_after); pf.space_before = Pt(0); pf.line_spacing = line
    if indent: pf.left_indent = Pt(indent)
    if hang: pf.first_line_indent = Pt(-hang)
    if text: _fmt_run(p.add_run(text), size, bold, italic)
    return p


def h1(doc, text):
    p = para(doc, text, size=12, bold=True, align=AL.CENTER, space_after=12, line=1.3)
    p.paragraph_format.space_before = Pt(6); return p


def h2(doc, text):
    p = para(doc, text, size=10, bold=True, align=AL.LEFT, space_after=4, line=1.3)
    p.paragraph_format.space_before = Pt(6); return p


def label(doc, text, size=9, align=AL.CENTER, bold=False, italic=False, after=0):
    return para(doc, text, size=size, bold=bold, italic=italic, align=align, space_after=after, line=1.3)


def table(doc, rows, caption):
    label(doc, caption, size=9, align=AL.CENTER, bold=True, after=2)
    t = doc.add_table(rows=len(rows), cols=len(rows[0])); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            cp = t.cell(i, j).paragraphs[0]; cp.alignment = AL.CENTER; cp.paragraph_format.line_spacing = 1.15
            _fmt_run(cp.add_run(str(cell)), 8.5, bold=(i == 0))
    para(doc, "", size=9, space_after=4); return t


def setup_section(sec, cols):
    sec.page_width = Mm(190); sec.page_height = Mm(260)
    sec.top_margin = Mm(20); sec.bottom_margin = Mm(20); sec.left_margin = Mm(20); sec.right_margin = Mm(20)
    sec.header_distance = Mm(15); sec.footer_distance = Mm(0)
    sectPr = sec._sectPr; c = sectPr.find(qn("w:cols"))
    if c is None: c = OxmlElement("w:cols"); sectPr.append(c)
    c.set(qn("w:num"), str(cols)); c.set(qn("w:space"), "360")


doc = Document()
st = doc.styles["Normal"]; st.font.name = BODY_FONT; st.font.size = Pt(9)
st.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT); st.paragraph_format.line_spacing = 1.6

# ── 전치(1단) ──
label(doc, "비례배분 기여도의 영-기여 흡수(zero-credit absorption)와", size=15, bold=True)
label(doc, "임계 가상비용에 의한 수동화 해소", size=15, bold=True)
label(doc, "— tabular Monte-Carlo 헤즈업 홀덤 학습의 사례연구 —", size=11, bold=True, after=8)
label(doc, "[저자명]*", size=10.5)
label(doc, "[소속기관] [직위]", size=9.5, after=8)
label(doc, "국문 요약", size=11, bold=True, after=4)
para(doc,
     "표 기반 Monte-Carlo(MC) 제어에서 각 행동의 보상을 그 행동의 투자액 비율로 배분하는 비례배분"
     "(proportional) 기여도 정형화는 분산을 줄이는 합리적 선택이나, 비용 0 행동(예: CHECK)의 credit이 매 "
     "에피소드 정확히 0이 되어 그 MC 고정점이 참값과 무관하게 0에 고정되는 구조적 병리 — 영-기여 "
     "흡수(zero-credit absorption, ZCA) — 를 남긴다. 본 연구는 이를 세 층위에서 다룬다. (이론) 최소 toy "
     "MDP에서 영-고정점과 그로 인한 양방향 오순위(참값이 음수인 체크의 과대평가 = 흡수, 참값이 양수인 "
     "체크의 과소평가 = 은폐)를 증명하고, 이 병리가 특정 배분식이 아니라 기여도-비례 배분 계열 전체"
     "(기여도 0 → credit 0)의 구조적 성질임을 보이며, 낙관적 초기화의 일시적 흡수와 질적으로 구별하고"
     "(영구·고정점≠참값), 해소에 필요한 임계 가상비용 ε_min을 유도한다. (실측) 2,048-셀 헤즈업 홀덤 "
     "봇에서 Q(CHECK)가 정확히 0으로 붕괴함을 관측하고(표준 MC는 −73~+120칩 분포), 그 정책적 결과가 "
     "수동화임을 행동 수준에서 보인다 — 턴에서 체크 65%의 소극 정책. (처방 실증) 1칩 가상비용은 회복을 "
     "재현하지 못하고 학습 seed에 따라 부호가 요동하지만(6-seed 평균 −117 mbb/g, 양수 2/6 — 전이구간), "
     "임계를 넘는 상수 5칩 이상이면 수동화가 풀리고(턴 소액 베팅 65%로 전환) 이 병리를 착취하던 미학습 "
     "상대(무작위 정책)에 대한 성능이 학습 seed 전반에서 회복된다(−318±123 → +1546±535 mbb/g, 0/6 → "
     "6/6 양수). 결정시점 상수 비용(fixed-K)·체크시점 팟 비례 등 인과적으로 깨끗한 변형에서 재현되어 "
     "사후정보·구현 인공물 가설은 통제실험으로 배제하였다. 끝으로 이 효과가 일반화가 아님을 정직하게 "
     "보고한다 — 사전 고정한 비용으로 4종 미학습 상대에 대해 검증한 결과 이득은 상대 의존적이며"
     "(콜링스테이션 상대는 소폭 손해), 효과는 단일 상대 학습 구성에 한정된다.", size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("주제어 : "), 9, bold=True)
_fmt_run(p.add_run("강화학습, Monte-Carlo 제어, 기여도 배분, 보상 형성, 영-기여 흡수, 고정점, 정책 수동화, "
                   "tabular Q-learning, 헤즈업 노리밋 홀덤"), 9)

label(doc, "Zero-Credit Absorption under Proportional Credit Assignment and the Resolution of",
      size=11.5, bold=True)
label(doc, "Policy Passivity by a Threshold Virtual Cost — A Case Study in Tabular MC Hold'em",
      size=11.5, bold=True, after=8)
label(doc, "[Author Name]*", size=10.5)
label(doc, "[Affiliation], [Position]", size=9.5, after=8)
label(doc, "ABSTRACT", size=11, bold=True, after=4)
para(doc,
     "In tabular Monte-Carlo (MC) control, assigning each action a share of the terminal payoff "
     "proportional to its invested chips—proportional credit assignment—is a reasonable "
     "variance-reduction choice, but it leaves a structural pathology for zero-cost actions (e.g., "
     "CHECK): their credit is exactly zero every episode, pinning the MC fixed point at zero "
     "regardless of true value — Zero-Credit Absorption (ZCA). We treat this at three levels. "
     "(Theory) In a minimal toy MDP we prove the zero fixed point and the resulting bidirectional "
     "mis-ranking (overvaluing bad checks = absorption; undervaluing good checks = masking), show "
     "that the pathology is a structural property of the whole family of contribution-proportional "
     "schemes (zero contribution → zero credit) rather than of one particular formula, distinguish "
     "it from the transient absorption of optimistic initialization (permanent; fixed point ≠ true "
     "value), and derive the threshold virtual cost ε_min required to dissolve it. (Measurement) In "
     "a 2,048-cell heads-up hold'em agent, the learned Q(CHECK) collapses exactly to zero (vs. −73 "
     "to +120 chips under standard MC), and we show behaviorally that the policy consequence is "
     "passivity — checking 65% of turn decisions. (Remedy) A one-chip cost fails to reproduce "
     "recovery, with its sign fluctuating across training seeds (mean −117 mbb/g, 2/6 positive — a "
     "transition zone), but a constant cost of five chips or more exceeds the threshold: passivity "
     "dissolves (65% small bets on the turn) and performance against the unseen opponent that "
     "exploited this pathology (a random policy) recovers across training seeds (−318±123 → "
     "+1546±535 mbb/g; 0/6 → 6/6 positive). The recovery reproduces under causally clean variants "
     "(decision-time constant cost; check-time pot fraction), ruling out hindsight-information and "
     "implementation-artifact explanations by controlled experiments. Finally, we honestly report "
     "that this is not generalization: with the cost fixed in advance, gains against four held-out "
     "opponents are opponent-dependent (a slight loss against a calling station), and the effect is "
     "confined to single-opponent training.", size=9, space_after=4)
p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6); p.paragraph_format.line_spacing = 1.6
_fmt_run(p.add_run("Keywords : "), 9, bold=True)
_fmt_run(p.add_run("reinforcement learning, Monte-Carlo control, credit assignment, reward shaping, "
                   "zero-credit absorption, fixed point, policy passivity, tabular Q-learning, "
                   "heads-up no-limit hold'em"), 9)

# ── 본문(2단) ──
body = doc.add_section(WD_SECTION.CONTINUOUS)

h1(doc, "1. 서론")
para(doc, "본 연구의 대상은 특정 게임이 아니라 구조적 기여도 배분이 남기는 측정 가능한 병리의 진단과, 그 "
          "병리를 해소하는 최소 개입의 임계 조건 규명이다. 표 기반 강화학습은 함수근사의 블랙박스 효과 "
          "없이 학습 동역학을 단일변수로 통제·관찰할 수 있으므로, 검증 환경으로는 헤즈업 노리밋 홀덤의 "
          "순수 tabular Monte-Carlo 학습을 선택한다. 이 선택은 네 가지 성질 때문이다: ① 비용 0 "
          "행동(CHECK)이 게임 규칙 자체에 내장되어 있어 인위적 삽입이 아니고, ② 보상이 말단에서 크고 "
          "확률적으로 실현되며, ③ 판돈 규모가 궤적마다 변해 임계 희석(4장)까지 시험되고, ④ 병리(수동화)를 "
          "착취하는 상대를 계측기로 쓸 수 있다.", space_after=6)
para(doc, "지연된 종단 보상을 궤적 위의 어느 행동에 얼마나 귀속시킬 것인가 — 기여도 배분(credit "
          "assignment; 국내 문헌에 따라 '신뢰할당'[20]) — 은 강화학습의 오랜 난제로, 최근 전용 조사 연구가 "
          "나올 만큼 활발히 연구되고 있다[19]. Monte-Carlo 제어에서 이 배분 방식은 학습 신호의 분산·편향을 "
          "좌우한다. 표준 Monte-Carlo는 편향이 없지만 분산이 크고, 자연스러운 대안은 각 행동이 판돈에 투입한 금액의 "
          "비율로 보상을 나누는 비례배분(proportional)이다. 그러나 비용 0 행동(투자액 0, 대표적으로 "
          "CHECK)의 비례 credit은 분자가 0이라 항상 0이고, 그 가치 추정의 고정점은 참값과 무관하게 0에 "
          "고정된다. 이 0은 행동 순위를 양방향으로 왜곡한다 — 참값이 음수인 체크는 과대평가되어 학습된 "
          "음수 행동들을 greedy에서 가리고(흡수), 참값이 양수인 체크는 과소평가되어 열등한 +EV 행동에 "
          "밀린다(은폐).", space_after=6)
para(doc, "본 논문의 기여는 다음과 같다. (1) 진단 — 위 현상을 영-기여 흡수(ZCA)로 명명해 toy MDP에서 "
          "양방향 오순위로 증명하고, 이것이 특정 배분식의 결함이 아니라 기여도-비례 배분 계열 전체의 "
          "구조적 성질(기여도 함수가 비용 0 행동에 0을 주는 모든 방식 — Lemma 3)임을 보이며(3장), 실제 "
          "2,048-셀 에이전트에서 Q(CHECK)의 정확한 영-고정점과 그 정책적 결과(턴 결정의 65%를 체크하는 "
          "수동화)를 측정한다(5장). (2) 임계 이론과 실증 — 비용 0 행동에 가상비용 ε을 부여하면 고정점이 "
          "참값 방향으로 풀리며, 해소에 필요한 임계 ε_min을 유도하고(4장), 실제 시스템에서 1칩은 회복을 "
          "재현하지 못하고 부호가 seed에 따라 요동하며(전이구간)·임계를 충분히 넘는 상수 5칩 이상은 전 "
          "seed 유효함을 재현한다(5장) — 이론이 예측한 임계의 존재가 단조 용량-반응으로 확인된 사례다"
          "(임계의 위치까지 예측하는 것은 아니다 — 6장 한계 ①). 아울러 동일 예산의 표준 MC와 3자 비교로 "
          "처방의 실용적 가치 — 저분산 배분을 보존한 채 병리만 제거 — 를 분리하고(5.3.1절), 표준 처방 "
          "후보(탐색 강화·낙관적 초기화·일률 벌점)가 전부 병리를 해소하지 못함을 통제 비교로 보인다"
          "(5.5절). (3) 정직한 범위 규정 — 이 회복이 \"일반화\"가 아니라 병리를 착취하던 상대에 대한 "
          "회복임을 사전 고정 비용의 홀드아웃 검증으로 보이고(콜링스테이션 상대는 소폭 손해), 평가 "
          "과정에서 발견·정정한 구현 인공물(credit 폴백 누수)과 단일 seed의 위험을 방법론적 부정적 결과로 "
          "보고한다(5·6장).", space_after=6)

h1(doc, "2. 관련 연구")
para(doc, "기존 알고리즘의 병리를 명명·진단하는 연구는 강화학습에서 확립된 계보를 갖는다: Q-learning의 "
          "과대추정 편향 진단[15]은 Double Q-learning[16]이라는 처방으로 이어졌고, 최근에도 primacy "
          "bias[17]·dormant neuron[18]처럼 병리 명명 → 진단 → 최소 처방의 형식이 반복된다. 본 연구는 "
          "같은 형식을 기여도 배분의 영-고정점에 적용한 것이다.", space_after=6)
para(doc, "기여도 배분을 직접 공략하는 해법 연구는 두 축으로 정리된다: 지연 보상을 어느 시점의 행동에 "
          "귀속시키는가(시간축 — 재분배·조사 연구[5][19])와, 공동 보상을 어느 에이전트에 귀속시키는가"
          "(에이전트축 — 반사실적 기준선[2]·가치 분해 계열). 국내 문헌에서도 이 문제는 다중 에이전트 맥락의 "
          "'신뢰할당'으로 소개되고[20] 가치 분해 계열 해법의 적용 연구가 이어지고 있으나[21], 배분식 자체가 "
          "특정 행동 부류에 남기는 병리를 다루는 원저는 확인되지 않는다. 본 연구는 그 답들 중 기여도-비례 "
          "계열이 배분식 자체에 심는 실패 모드를 다루며, 최신 조사 연구[19]의 실패 모드 분류에도 본 병리"
          "(비용 0 행동의 영-고정점)는 명명되어 있지 않다.", space_after=6)
para(doc, "ZCA를 이루는 부품은 인접 문헌에 모두 존재하나, 본 연구는 그 부호를 뒤집어 같은 구조를 실패 "
          "모드로 진단한다. 협력게임 이론의 Shapley value[3]·Shapley Q-value[4]는 null-player 공리(한계 "
          "기여 0 → 배분 0)를 공정성의 바람직한 성질로 둔다 — 본 연구의 \"비용 0 → credit 0\"은 그 "
          "대응물이나 진단 방향이 반대다. 같은 직관은 difference rewards[1]·COMA[2]에도 깔려 있다. "
          "RUDDER[5]는 return-equivalent 재분배가 최적 정책을 보존함을 보이는데, 비례배분은 비용 0 행동에 "
          "대해 return-equivalent가 아니며 ZCA는 그 비등가 영역에서 발생한다(보존 정리의 대우). 처방 "
          "측면에서 가상비용은 PBRS가 Q-value 초기화와 등가[6][7]인 틀 안에 위치하나, \"비용 0 행동에 "
          "선택적으로 가상비용을 부여해 영-고정 오순위를 임계 조건 하에 해소한다\"는 처방과 그 임계의 "
          "유도·실증은 확인되지 않았고, 가장 가까운 action-penalty[8]는 모든 행동에 일률적이라 선택적이지 "
          "않다. 미방문/낙관적 0의 일시적 흡수[9]는 초기화 기인으로, 비례 credit이 구조적으로 0을 "
          "재고정하는 본 연구와 기제가 다르다. 단일 seed의 오도 위험은 재현성 문헌[10][11][12]이 "
          "확립했으며, 본 연구는 모든 성능 주장을 6개 학습 seed에서 검증하고 단일 seed가 실제로 오도했던 "
          "사례를 6장에서 보고한다.", space_after=6)

h1(doc, "3. 영-기여 흡수의 형식적 특성화 — 양방향 오순위")
para(doc, "문제 설정(toy MDP). 결정 상태 s에서 행동 a₁을 한 번 선택한다. 각 행동은 투자액 inv(a₁)≥0을 "
          "가지며 비용 0 행동(CHECK)은 inv=0이다. 선택 후 궤적은 투자액 c>0인 후속 행동을 포함하고 종단 "
          "보상 P로 끝난다(γ=1). μ(a₁):=E[P|a₁]로 두면 참값은 q*(s,a₁)=μ(a₁)이다.", space_after=4)
table(doc, [["방식 (scheme)", "R(s, a₁)"],
            ["표준 MC", "P"],
            ["비례배분 (PROP)", "[ inv(a₁) / (inv(a₁)+c) ] · P"],
            ["가상비용 (VIC)", "PROP과 동일, 단 inv(CHECK) ← ε > 0"]],
      "<Table 1> 행동별 return R(s, a₁) / Per-action return")
para(doc, "Lemma 1 (표준 MC). Q_std(s,a₁) → μ(a₁) = q*(s,a₁).", space_after=4)
para(doc, "Lemma 2 (영-고정점). Q_prop(s,CHECK)=[0/(0+c)]·μ(CHECK)=0 (∀ μ(CHECK)). 참값과 무관하고, "
          "R≡0이라 표본 분산도 0인 구조적 고정점이다(초기화 잔재 아님 — 방문할수록 0으로 재고정).", space_after=4)
para(doc, "Lemma 3 (계열 일반화). 임의의 기여도 함수 φ≥0에 대해 credit(a)=[φ(a)/Σφ]·P 꼴의 비례 "
          "credit은 φ(a)=0인 행동에서 Lemma 2와 동일한 영-고정점을 갖는다 — ZCA ⟺ φ(비용 0 행동)=0. "
          "투자액(φ=inv)뿐 아니라 칩과 무관한 공격성 지표(φ=베팅 여부 1/0)에서도 재현되고, 균등 기여"
          "(φ≡1)와 표준 MC(return-equivalent)는 면역이다. 즉 ZCA는 본 연구가 채택한 특정 배분식의 결함이 "
          "아니라, null-player 공리(기여 0 → 배분 0)[3]를 가치 학습 신호로 쓰는 계열 전체의 구조적 "
          "함정이다.", space_after=4)
para(doc, "Theorem (흡수 — 나쁜 체크의 과대평가). μ_C < μ_B < 0이면 비례배분 greedy는 Q_prop(CHECK)=0 > "
          "[b/(b+c)]μ_B이므로 열등한 CHECK를 선택한다(표준 MC는 최적 BET 선택).", space_after=4)
para(doc, "Theorem′ (은폐 — 좋은 체크의 과소평가). μ_C > μ_B > 0이면 비례배분 greedy는 Q_prop(CHECK)=0 < "
          "[b/(b+c)]μ_B이므로 참-최선인 CHECK를 버리고 열등한 +EV BET을 선택한다(트랩·팟컨트롤 라인의 "
          "상실). 즉 ZCA는 특정 부호 조건의 문제가 아니라 CHECK의 참값이 0이 아닌 모든 결정 상태에서의 "
          "양방향 오순위다.", space_after=4)
para(doc, "Proposition 1 (낙관적 초기화와의 구분). 0-초기화의 0-선호는 충분한 방문 후 Q→μ로 소거되는 "
          "일시적 현상이고 고정점은 참값이다(Lemma 1). 반면 Q_prop(CHECK)=0은 수렴 후에도 유지되는 "
          "고정점이며 참값과 다르다(Lemma 2). 낙관적 초기화는 \"아직 학습 못 해서\", ZCA는 \"구조적으로 "
          "학습 불가라서\"다.", space_after=6)

h1(doc, "4. 임계 가상비용")
para(doc, "Proposition 2 (해소 임계). CHECK에 가상 투자 ε>0을 주면 Q_vic(CHECK)=[ε/(ε+c)]·μ_C로 고정점이 "
          "참값 방향으로 풀린다. 두 모드 공통으로, greedy가 참-우월 행동을 선택할 필요충분조건은 "
          "k:=[b/(b+c)](μ_B/μ_C)∈(0,1)에 대해  ε > ε_min = k·c/(1−k)  이다(흡수·은폐 모드에서 동형; 수치 "
          "검증 일치). 임계 미달이면 Q_vic(CHECK)≈0으로 오순위가 잔존한다 — 이 예측이 5장에서 \"1칩은 "
          "회복 재현 실패(부호 요동), 5칩 이상 전 seed 유효\"로 확인된다. 팟 규모 c가 궤적마다 변하는 "
          "환경에서는 고정 소액 ε이 큰 팟에서 희석되므로(ε/(ε+c)→0), 임계는 대표 팟 규모 기준으로 "
          "설정해야 한다. 나아가 결정 상태(셀)마다 팟 규모가 달라 ε_min도 셀마다 다르므로, 고정 ε이 일부 "
          "셀에서만 임계를 넘는 중간 크기에서는 부분적·seed 의존적 회복(전이구간)이 예상된다 — 5장에서 "
          "1칩이 정확히 이 양상을 보인다.", space_after=6)
para(doc, "탈출 메커니즘 비교. 무정보 처방은 이 오순위를 풀지 못한다 — 결정시 노이즈는 고정점 0을 바꾸지 "
          "못하고, 동률 처리(tie-break)는 strict 부등식(0 > 음수)에 미발동이다. 가상비용과 고정 페널티만 "
          "고정점을 옮기며, 가상비용의 고정점만 참값 μ_C에 연동(informed)된다.", space_after=6)

h1(doc, "5. 실제 MDP에서의 측정과 처방 실증")
para(doc, "환경: (라운드 × 포지션 × 핸드 버킷 × 직전 행동) 2,048-셀 추상화, 8개 행동, softmax 탐색(온도 "
          "감쇠), 2M 에피소드, 학습 상대 TAG(타이트-공격 룰 정책) 단일. 평가는 100k 게임 × 5회(평가 seed "
          "고정) raw greedy, 성능 주장은 학습 seed 6개(1–5, 42)에서 검증. 단위 mbb/g(=payoff×500). 모든 "
          "학습은 6장에서 보고하는 credit 폴백 인공물을 제거(clean)한 코드로 수행했다.", space_after=4)

h2(doc, "5.1 영-고정점의 실측")
para(doc, "표준 MC(PURE)로 학습하면 Q(CHECK)가 −72.9~+119.8칩으로 펼쳐지지만(평균 |Q(CHECK)| 4.48), "
          "비례배분에서는 모든 셀에서 정확히 0이다(평균 0.00, 범위 [0,0]) — Lemma 2의 직접 관측이며 학습 "
          "seed·표본과 무관한 구조적 사실이다. 이 불변식(\"PROP+무개입 → Q(CHECK)==0.000\")은 자동 "
          "테스트로 저장소에 고정하였다.", space_after=4)

h2(doc, "5.2 병리의 정책적 결과 — 수동화")
para(doc, "영-고정점의 행동적 결과를 보기 위해 학습된 정책의 행동 분포를 측정하면(20k 게임), 무개입 정책은 "
          "턴 결정의 65%를 체크한다(<Figure 3>). 공격 행동의 Q가 0 근방(약학습)이거나 음수인 셀에서 "
          "CHECK=0이 greedy를 차지하기 때문이다(Theorem의 실제 형태). 이 수동화는 그것을 착취하는 "
          "상대에게 체계적으로 진다 — 무작위 정책(미학습 상대) 대상 성능이 6개 학습 seed 전부에서 "
          "음수다(−318±123 mbb/g, 0/6 양수). 무작위 상대조차 이기지 못하는 것은 강함의 문제가 아니라 "
          "병리의 증거다(단순 공격 룰 정책은 같은 상대에 +11,705를 얻는다).", space_after=4)

h2(doc, "5.3 임계 실증 — 1칩은 전이구간, 5칩 이상은 유효")
para(doc, "가상비용의 크기를 단일변수로 스윕한 결과(<Table 2>, <Figure 1>), Proposition 2의 예측과 "
          "정합하게 용량-반응이 단조하며, 임계를 충분히 넘는 개입에서만 전 seed 양수가 된다.", space_after=4)
table(doc, [["개입 (CHECK 가상 invest)", "vs 무작위 (mbb/g)", "양수 seed", "판정"],
            ["(참조) 표준 MC — 배분 제거", "−312±91", "0/5", "병리 없음·양 지표 열세(5.3.1)"],
            ["없음 (ε=0)", "−318±123", "0/6", "수동화(기준선)"],
            ["상수 1칩", "−117±252", "2/6¹", "전이구간 — 부호 요동"],
            ["상수 5칩", "+1230±693", "5/5", "임계 초과 — 유효"],
            ["상수 20칩", "+1659±825", "5/5", "유효"],
            ["상수 60칩", "+806±451", "5/5", "유효(과용량 방향)"],
            ["체크시점 팟 30%", "+1546±535", "6/6", "유효"],
            ["체크시점 팟 ≤15%", "−276 ~ −94", "0~3/5", "전이구간 — 비신뢰"]],
      "<Table 2> 가상비용 크기별 성능 / Performance by virtual-cost magnitude (학습 seed 5–6개)")
para(doc, "¹ 1칩 조건의 초기 단일 seed(42) 측정은 −376으로 \"무효\"를 시사했으나, 6-seed 확장에서 부호 "
          "요동(2/6 양수)이 드러났다 — 단일 seed의 오도 위험(6장)의 추가 사례이며, 셀별 임계 이질성에 "
          "의한 전이구간 해석(4장)과 정합한다. 행동 수준에서 임계 초과 개입은 수동화를 정확히 뒤집는다"
          "(<Figure 3>) — 턴 체크 65% → 턴 소액 베팅 65%, 승률 51.8→54.5%, 평균 승리 팟 +19.9→+25.3칩. "
          "즉 Q(CHECK)가 참값(무작위 상대에게 턴 체크는 지분 포기 = 음수)을 학습하자 greedy가 베팅으로 "
          "전환한 것이다. 결정시점 상수 비용(fixed-K)에서 재현되므로 사후정보(최종 팟 사용) 가설은 "
          "배제되고, 이론(결정시점 상수 ε)과 정확히 대응한다.", space_after=4)

h2(doc, "5.3.1 표준 MC와의 동일 예산 비교 — \"비례배분을 왜 쓰는가\"")
para(doc, "\"비례배분이 병리를 만드니 처음부터 표준 MC를 쓰면 된다\"는 자연스러운 반박을 동일 조건 3자 "
          "비교로 검증하였다(single-TAG·2M·seed 5개). 표준 MC(PURE)는 vs TAG +115±25, vs 무작위 "
          "−312±91(0/5 양수)로 양 지표 모두 열세다 — 비례배분+임계 비용(+909 / +1230~+1546) 대비 크게 "
          "낮다. 학습 곡선(<Figure 2>)은 그 원인이 학습 속도가 아니라 안정성임을 보여준다: 표준 MC는 학습 "
          "중반까지 비례배분과 대등한 성능(vs TAG 약 +900)을 보이다가 탐색 온도가 낮아지는 후반(약 1.4M "
          "에피소드)에 붕괴한다(고분산 MC의 알려진 불안정). 비례배분 계열은 개입 유무와 무관하게 붕괴하지 "
          "않는다. 즉 비례배분은 저분산으로 후반 안정성을 사고, 그 대가인 ZCA를 임계 비용이 치료한다 — "
          "세 방식 중 \"강하게 배우고 병도 없는\" 조합은 비례배분+임계 비용뿐이다.", space_after=4)

h2(doc, "5.4 이것은 일반화가 아니다 — 상대별 손익과 범위")
para(doc, "비용을 사전에 고정하고(체크시점 팟 30%) 학습에 쓰지 않은 4종 상대로 검증하면(<Table 3>), "
          "이득은 상대 의존적이다.", space_after=4)
table(doc, [["홀드아웃 상대", "Δ(개입−무개입)", "개선 seed"],
            ["무작위 (병리 착취 상대)", "+1,864", "6/6"],
            ["LAG (루즈-공격)", "+73", "3/5"],
            ["Maniac (초공격)", "+307", "3/5"],
            ["Nit (초타이트)", "+16", "5/5"],
            ["Station (콜링스테이션)", "−177", "0/5"]],
      "<Table 3> 홀드아웃 상대별 손익 / Held-out opponent deltas (mbb/g, seed 5개 평균)")
para(doc, "수동화 해소는 스타일 변화(공격성 증가)이므로, 공격이 통하는 상대(무작위: 큰 회복)와 통하지 않는 "
          "상대(콜링스테이션: 베팅에 폴드하지 않으므로 소폭 손해)에서 손익이 갈린다. 따라서 본 효과의 "
          "정확한 서술은 \"분포 외 일반화\"가 아니라 \"영-고정점이 만든 수동화 병리의 해소 — 그 병리를 "
          "착취하던 상대에 대한 회복\"이다. 또한 효과는 단일 상대(TAG) 학습에 한정된다 — 5종 상대 "
          "순환/혼합 학습에서는 같은 개입으로 무작위 상대 성능이 회복되지 않았다(−112/−159). 학습 분포 내 "
          "성능(vs TAG)은 모든 조건에서 +849~+952로 보존된다.", space_after=6)

h2(doc, "5.5 경쟁 처방 비교 — 표준 처방 후보는 병리를 해소하지 못한다")
para(doc, "4장의 이론적 기각(무정보 처방은 고정점을 옮기지 못함)을 동일 조건 통제 비교로 실증하였다(각 "
          "학습 seed 5개, 그 외 규약 동일). 비교 대상은 ZCA에 대해 제안될 법한 표준 처방 세 가지다: ① "
          "탐색 강화(softmax 온도 하한 2.0 — \"더 탐색하면 배우지 않을까\"), ② 낙관적 초기화(전 셀 "
          "Q₀=+50 — Proposition 1의 검증), ③ 일률 가산 벌점(전 행동의 credit에서 5칩 차감 — "
          "action-penalty[8]류의 비선택적 비용).", space_after=4)
table(doc, [["처방", "vs 무작위", "양수 seed", "vs TAG"],
            ["없음 (기준선)", "−318±123", "0/6", "+866"],
            ["① 탐색 강화 (온도≥2.0)", "−291±267", "1/5", "+880±25 (보존)"],
            ["② 낙관적 초기화 (Q₀=50)", "−241±302", "1/5", "+184±43 (훼손)"],
            ["③ 일률 벌점 (전 행동 −5칩)", "−485±32", "0/5", "+119±34 (훼손)"],
            ["선택적 가상비용 (CHECK만 5칩)", "+1230±693²", "5/5", "+909±50 (보존)"]],
      "<Table 4> 경쟁 처방 비교 / Rival remedies (mbb/g, 학습 seed 5–6개)")
para(doc, "² 5칩은 (i) ③ 일률 벌점과 크기를 일치시켜 선택성만을 분리하고, (ii) 임계를 넘는 최소 유효 "
          "용량이라는 보수적 대표값이다. 더 높은 수치의 조건(체크시점 팟 30% +1546±535·6/6, 상수 20칩 "
          "+1659±825·5/5)으로 바꿔도 결론은 불변하며, 조건 간 순위는 표준편차 중첩으로 주장하지 않는다"
          "(<Table 2>).", space_after=4)
para(doc, "①은 예측대로 무효다 — 방문이 늘수록 0으로 재고정되므로 탐색량은 고정점을 바꾸지 못한다(학습 "
          "분포 내 성능은 보존). ②는 회복 실패에 더해 학습 분포 내 성능까지 훼손했다 — 낙관값이 전 "
          "셀에서 씻겨 내려가는 데 학습 예산이 소모된다(Proposition 1의 일시성이 유한 예산에서는 비용이 "
          "됨). ③은 이론 예측(순위 불변 → 기준선과 동일)보다 나빴다: 고정점의 순위는 불변이나, 벌점이 "
          "방문한 셀만 −5칩 이동시켜 미방문 셀의 0-초기값과 상대 격차를 만들고(의도치 않은 상대적 낙관 "
          "초기화), 유한 예산 학습을 양 지표 모두에서 훼손했다 — 이 예측 편차 자체를 부정적 결과로 함께 "
          "보고한다. 다섯 조건 중 양 지표를 지키며 병리를 해소한 것은 선택적 임계 가상비용뿐이며, 특히 "
          "③과의 대비는 처방의 본질이 \"비용 부여\" 일반이 아니라 비용 0 행동에 대한 선택성임을 분리 "
          "실증한다.", space_after=6)

h1(doc, "6. 방법론적 부정적 결과와 한계 (정직)")
para(doc, "(1) credit 폴백 인공물의 발견과 정정. 초기 구현은 총투자 0인 핸드(전부 체크/폴드)에서 비례 "
          "분배가 미정의라 균등분배(payoff/n)로 폴백했고, 이 누수가 비용 0 행동에 비(非)비례 credit을 "
          "흘렸다. 단일 학습 seed(42)와 결합해 \"1칩 가상비용이 분포 외 일반화의 필요조건\"이라는 허위 "
          "신호를 만들었으며, 폴백 격리런(원본 수치 완전 재현)과 제거 재실행(효과 소멸)으로 인공물임을 "
          "확정하고 폐기하였다 — 단일 seed 오도 위험[10][11][12]의 실제 사례. 본 논문의 모든 수치는 폴백 "
          "제거 후의 것이며, 나아가 개입을 실투자 핸드에만 적용하는 격리 실험으로 5장의 회복이 폴백 유사 "
          "신호가 아니라 CHECK credit 경로임을 확인했다(실투자-한정 +1245(5/5) ≈ 전체 적용 +1546; "
          "폴백-신호-한정 +206·고분산).", space_after=4)
para(doc, "(2) 한계. ① 임계의 정확한 위치는 미확정 — 전이구간(팟 8~15%·상수 1칩 부근)은 seed 요동이 "
          "크며(1칩 6-seed 평균 −117·2/6 양수), \"충분 초과 시 전 seed 유효\"라는 임계의 존재만 주장한다. "
          "② 효과의 범위는 단일 상대 학습·단일 게임(HUNL)·단일 추상화에 한정된다. ③ 홀드아웃 상대도 자작 "
          "룰 정책이며 외부·균형 상대(예: 근사 GTO)에 대한 검증은 향후 과제다. ④ ZCA의 부품(null-player "
          "공리·return-equivalence·낙관적 초기화·PBRS≡Q-init)은 기지이며, 기여는 이들을 실패 모드로 "
          "재진단하고 임계를 유도·실증하며 범위를 정직히 규정한 데 있다.", space_after=6)

h1(doc, "7. 결론 및 향후 연구")
para(doc, "비례배분 기여도는 비용 0 행동에 구조적 영-고정점(ZCA)을 남기고, 이는 행동 순위를 양방향으로 "
          "왜곡하며(흡수·은폐), 낙관적 초기화의 일시적 0-선호와 달리 방문으로 해소되지 않는다 — 이상을 "
          "toy MDP에서 증명하고 실제 2,048-셀 에이전트에서 Q(CHECK)=0의 정확한 붕괴와 그 결과인 정책 "
          "수동화(턴 체크 65%)로 실측하였다. 유도한 임계 조건의 예측은 단조 용량-반응으로 확인되었다: "
          "무개입은 전 seed 음수(0/6), 1칩은 부호가 요동하는 전이구간(2/6), 상수 5칩 이상이면 수동화가 "
          "풀리고(턴 소액 베팅 65%) 병리를 착취하던 미학습 상대에 대한 성능이 학습 seed 전반에서 "
          "회복된다(5/5·6/6 양수). 확인된 것은 임계의 존재이며 위치의 예측이 아니다(6장 한계 ①). "
          "사후정보·구현 인공물·평가 순환성 가설은 각각 결정시점 상수 비용, 격리 실험, 사전 고정 "
          "홀드아웃으로 배제하였고, \"표준 MC를 쓰면 된다\"는 대안은 동일 예산 3자 비교(양 지표 열세)로 "
          "기각하였다. 동시에 이 회복이 일반화가 아니라 상대 의존적 스타일 변화임을 — 콜링스테이션 상대의 "
          "소폭 손해까지 — 정직하게 보고하였다. 본 사례연구의 가치는, 구조적 기여도 배분의 병리를 "
          "증명·측정하고, 처방의 작동 조건(임계)을 이론-실험 대응으로 규명하며, 그 한계를 부정적 결과와 "
          "함께 드러낸 데 있다.", space_after=4)
para(doc, "향후 연구. ZCA의 발생 조건은 게임 특수적이지 않다 — ① 성과가 말단에 한 번 실현되고, ② 행동이 "
          "성과와 같은 단위의 자원 투입이며, ③ 투입 0 행동이 존재하는 순차 결정 환경이면 계열 정리(Lemma "
          "3)가 동일 병리의 발생을 예측한다. 이 세 조건은 예컨대 V2G 충·방전 스케줄링(대기 행동의 투입 0, "
          "청구 주기 말 요금)과 에너지 하베스팅 IoT 전송 스케줄링(sleep 행동의 투입 0, 보고 주기 말 "
          "효용)에서 성립한다. 두 환경에서의 검증 — 특히 요금 구조가 균질한 V2G에서의 임계 위치 사전 "
          "계산, sleep의 대기전력이라는 실재 비용에 의한 가상비용의 물리적 해석 — 을 후속 연구로 "
          "진행한다.", space_after=8)

# ── 그림 (초안 배치: 말미 1단 섹션) ──
figsec = doc.add_section(WD_SECTION.CONTINUOUS)
FIGS = [
    ("figs/fig1_dose_response.png",
     "<Figure 1> Performance vs. virtual-cost magnitude (per-seed dots, mean±SD)"),
    ("figs/fig2_learning_curves.png",
     "<Figure 2> Learning curves vs. training opponent (5-seed mean; standard MC collapses late)"),
    ("figs/fig3_turn_behavior.png",
     "<Figure 3> Turn action distribution: passivity (65% check) flips to small bets under threshold cost"),
]
for path, cap in FIGS:
    p = doc.add_paragraph(); p.alignment = AL.CENTER
    p.add_run().add_picture(path, width=Mm(140))
    label(doc, cap, size=8.5, align=AL.CENTER, after=8)

refsec = doc.add_section(WD_SECTION.CONTINUOUS)
para(doc, "참고문헌", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
refs = [
    '[1] D. H. Wolpert and K. Tumer, "Optimal payoff functions for members of collectives," Advances in Complex Systems, Vol. 4, No. 2/3, pp. 265-279, 2001.',
    '[2] J. Foerster, G. Farquhar, T. Afouras, N. Nardelli, and S. Whiteson, "Counterfactual multi-agent policy gradients," Proc. AAAI Conf. on Artificial Intelligence, pp. 2974-2982, 2018.',
    '[3] L. S. Shapley, "A value for n-person games," Contributions to the Theory of Games II, Princeton Univ. Press, pp. 307-317, 1953.',
    '[4] J. Wang, Y. Zhang, T.-K. Kim, and Y. Gu, "Shapley Q-value: A local reward approach to solve global reward games," Proc. AAAI Conf. on Artificial Intelligence, pp. 7285-7292, 2020.',
    '[5] J. A. Arjona-Medina, M. Gillhofer, M. Widrich, T. Unterthiner, J. Brandstetter, and S. Hochreiter, "RUDDER: Return decomposition for delayed rewards," Advances in Neural Information Processing Systems, pp. 13544-13555, 2019.',
    '[6] A. Y. Ng, D. Harada, and S. Russell, "Policy invariance under reward transformations: Theory and application to reward shaping," Proc. Int. Conf. on Machine Learning, pp. 278-287, 1999.',
    '[7] E. Wiewiora, "Potential-based shaping and Q-value initialization are equivalent," Journal of Artificial Intelligence Research, Vol. 19, pp. 205-208, 2003.',
    '[8] S. Koenig and R. G. Simmons, "The effect of representation and knowledge on goal-directed exploration with reinforcement-learning algorithms," Machine Learning, Vol. 22, pp. 227-250, 1996.',
    '[9] T. Rashid, B. Peng, W. Boehmer, and S. Whiteson, "Optimistic exploration even with a pessimistic initialisation," Proc. Int. Conf. on Learning Representations, 2020.',
    '[10] P. Henderson, R. Islam, P. Bachman, J. Pineau, D. Precup, and D. Meger, "Deep reinforcement learning that matters," Proc. AAAI Conf. on Artificial Intelligence, pp. 3207-3214, 2018.',
    '[11] C. Colas, O. Sigaud, and P.-Y. Oudeyer, "How many random seeds? Statistical power analysis in deep reinforcement learning experiments," arXiv:1806.08295, 2018.',
    '[12] R. Agarwal, M. Schwarzer, P. S. Castro, A. C. Courville, and M. G. Bellemare, "Deep reinforcement learning at the edge of the statistical precipice," Advances in Neural Information Processing Systems, pp. 29304-29320, 2021.',
    '[13] J. Kim, "PokerKit: A comprehensive Python library for fine-grained multi-variant poker game simulations," IEEE Trans. on Games, 2023.',
    '[14] R. S. Sutton and A. G. Barto, Reinforcement Learning: An Introduction, 2nd ed., MIT Press, 2018.',
    '[15] S. Thrun and A. Schwartz, "Issues in using function approximation for reinforcement learning," Proc. of the 1993 Connectionist Models Summer School, pp. 255-263, 1993.',
    '[16] H. van Hasselt, "Double Q-learning," Advances in Neural Information Processing Systems, pp. 2613-2621, 2010.',
    '[17] E. Nikishin, M. Schwarzer, P. D\'Oro, P.-L. Bacon, and A. Courville, "The primacy bias in deep reinforcement learning," Proc. Int. Conf. on Machine Learning, pp. 16828-16847, 2022.',
    '[18] G. Sokar, R. Agarwal, P. S. Castro, and U. Evci, "The dormant neuron phenomenon in deep reinforcement learning," Proc. Int. Conf. on Machine Learning, pp. 32145-32168, 2023.',
    '[19] E. Pignatelli, J. Ferret, M. Geist, T. Mesnard, H. van Hasselt, and L. Toni, "A survey of temporal credit assignment in deep reinforcement learning," arXiv:2312.01072, 2023.',
    '[20] 유병현, 데브라니 데비, 김현우, 송화전, 박경문, 이성원, "멀티 에이전트 강화학습 기술 동향," 전자통신동향분석, 제35권 제6호, pp. 137-149, 2020.',
    '[21] 김민경, "Ray RLlib 기반 QMIX와 RND를 이용한 희소 보상 전장 환경에서의 멀티에이전트 강화학습 협업," 한국컴퓨터정보학회논문지, 제29권 제1호, pp. 11-19, 2024.',
]
for r in refs:
    para(doc, r, size=9, align=AL.JUSTIFY, space_after=2, line=1.3, indent=16, hang=16)

para(doc, "저자소개", size=12, bold=True, align=AL.CENTER, space_after=6, line=1.3)
para(doc, "[저자명] (Author Name)  [정회원]", size=9, space_after=2)
para(doc, "[학사/석사/박사 취득연월 및 학위명]", size=9, space_after=2)
para(doc, "[현 소속기관 및 직위]  E-mail : [e-mail]", size=9, space_after=2)
para(doc, "관심분야 : 강화학습, 게임 AI, 학습 동역학 분석", size=9, space_after=2)

setup_section(doc.sections[0], cols=1)
setup_section(body, cols=2)
setup_section(figsec, cols=1)
setup_section(refsec, cols=2)
doc.save("zca_vic_논문초안_v3.docx")
print("saved: zca_vic_논문초안_v3.docx  (v3.1 동기화: Lemma3·전이구간·5.3.1·5.5·향후연구·[15-18])")
